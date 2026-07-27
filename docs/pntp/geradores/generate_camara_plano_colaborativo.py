from __future__ import annotations

import csv
import re
from collections import Counter, OrderedDict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
CSV_PATH = ROOT / "docs" / "pntp" / "criterios-pntp-2026.csv"
OUT_DIR = ROOT / "docs" / "pntp" / "solicitacoes-portal"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_PATH = OUT_DIR / "Plano_Colaborativo_Melhorias_Portal_Transparencia_Camara_Varginha.docx"

# Preset: compact_reference_guide.
# Named geometry override: A4 Brasil, 2 cm margins, 17 cm usable width.
# Named table override: dense_annex_table, 8.2 pt body for the 83-item PNTP checklist.
NAVY = "17324D"
BLUE = "1F5D8F"
LIGHT_BLUE = "EAF2F8"
PALE_BLUE = "F4F7FA"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "687581"
INK = "20262C"
WHITE = "FFFFFF"
RED = "A33A32"
PALE_RED = "FBEDEC"
GREEN = "297A5B"
PALE_GREEN = "EAF5EF"
GOLD = "8A6500"
PALE_GOLD = "FFF7DD"

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_CM = 2.0
CONTENT_CM = 17.0
CONTENT_DXA = 9638
TABLE_INDENT_DXA = 120

PORTAL_URL = "https://camaravarginha.mg.gov.br/transparencia"
PNTP_DOWNLOADS_URL = "https://radardatransparencia.atricon.org.br/downloads.html"
PNTP_CARTILHA_URL = "https://radardatransparencia.atricon.org.br/pdf/Cartilha-PNTP-2025.pdf"
RESULT_URL = "https://radartransparente.com.br/avaliacao/5518baf1-b161-4f79-b364-94c64e2af7bb"
LAI_URL = "https://www.planalto.gov.br/ccivil_03/_ato2011-2014/2011/lei/l12527.htm"
LOCAL_LAI_URL = (
    "https://sapl.varginha.mg.leg.br/media/sapl/public/normajuridica/"
    "2013/5880/lei5729-2013.pdf"
)
LRF_URL = "https://www.planalto.gov.br/ccivil_03/leis/lcp/lcp101.htm"
SIAFIC_URL = "https://www.planalto.gov.br/ccivil_03/_ato2019-2022/2020/decreto/d10540.htm"
PROCUREMENT_URL = "https://www.planalto.gov.br/ccivil_03/_ato2019-2022/2021/lei/l14133.htm"
OUVIDORIA_URL = "https://www.planalto.gov.br/ccivil_03/_ato2015-2018/2017/lei/l13460.htm"
LGPD_URL = "https://www.planalto.gov.br/ccivil_03/_ato2015-2018/2018/lei/l13709compilado.htm"
LBI_URL = "https://www.planalto.gov.br/ccivil_03/_ato2015-2018/2015/lei/l13146.htm"


def set_run_font(run, size: float, color: str = INK, bold=None, italic=None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **edges) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, attrs in edges.items():
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        for key, value in attrs.items():
            edge.set(qn(f"w:{key}"), str(value))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    tr_pr.append(node)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def set_repeat_keep(paragraph, keep_with_next=False) -> None:
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.widow_control = True
    if keep_with_next:
        paragraph.paragraph_format.keep_with_next = True


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=TABLE_INDENT_DXA) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_hyperlink(paragraph, text: str, url: str, color=BLUE) -> None:
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_color = OxmlElement("w:color")
    r_color.set(qn("w:val"), color)
    r_pr.append(r_color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "20")
    r_pr.append(size)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr, fld_char_2])
    set_run_font(run, 8.5, MID_GRAY)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_CM)
    section.bottom_margin = Cm(MARGIN_CM)
    section.left_margin = Cm(MARGIN_CM)
    section.right_margin = Cm(MARGIN_CM)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 11.5, NAVY, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Cm(CONTENT_CM))
    set_table_geometry(table, [6100, 3538], indent_dxa=0)
    table.cell(0, 0).text = "FISCALIZA VARGINHA"
    table.cell(0, 1).text = "CÂMARA MUNICIPAL"
    for idx, cell in enumerate(table.rows[0].cells):
        set_cell_shading(cell, NAVY if idx == 0 else BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            set_run_font(run, 9, WHITE, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run("Contribuição cidadã independente • Plano colaborativo • Página ")
    set_run_font(run, 8.5, MID_GRAY)
    add_page_field(p)


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("PLANO COLABORATIVO DE MELHORIA")
    set_run_font(run, 10, BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    set_repeat_keep(p, keep_with_next=True)
    run = p.add_run("Portal da Transparência da Câmara Municipal de Varginha")
    set_run_font(run, 24, NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(
        "Diagnóstico de falhas e plano de regularização integral em 60 dias"
    )
    set_run_font(run, 12.5, MID_GRAY, italic=True)

    table = doc.add_table(rows=4, cols=2)
    rows = [
        ("DESTINATÁRIO", "Mesa Diretora, Controladoria Interna e setores responsáveis"),
        ("PROPONENTE", "Fiscaliza Varginha - iniciativa cidadã independente"),
        ("DATA-BASE", "25 de julho de 2026"),
        ("PROPÓSITO", "Resolver as pendências e concluir as melhorias aplicáveis em até 60 dias"),
    ]
    for idx, (label, value) in enumerate(rows):
        table.cell(idx, 0).text = label
        table.cell(idx, 1).text = value
    set_table_geometry(table, [2100, 7538])
    for row_idx, row in enumerate(table.rows):
        prevent_row_split(row)
        for col_idx, cell in enumerate(row.cells):
            if col_idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                set_run_font(run, 9.3, NAVY if col_idx == 0 else INK, bold=col_idx == 0)
            set_cell_border(cell, bottom={"val": "single", "sz": "3", "color": "D8DEE3"})

    add_callout(
        doc,
        "Tom e finalidade",
        "Este documento reconhece os avanços já existentes, aponta falhas e atrasos que precisam "
        "ser enfrentados e propõe uma agenda prática de cooperação para conclusão em 60 dias. "
        "Não presume irregularidade nem substitui a avaliação oficial do Tribunal de Contas.",
        fill=PALE_BLUE,
        accent=BLUE,
    )


def add_body(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_repeat_keep(p)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, 10.5, INK, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, 10.5, INK)
    else:
        r = p.add_run(text)
        set_run_font(r, 10.5, INK)


def add_callout(doc: Document, title: str, body: str, fill=LIGHT_BLUE, accent=BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    prevent_row_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, left={"val": "single", "sz": "18", "color": accent})
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{title}. ")
    set_run_font(r, 10, NAVY, bold=True)
    r = p.add_run(body)
    set_run_font(r, 10, INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.space_before = Pt(0)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[int],
    font_size: float = 9.0,
    first_col_bold: bool = False,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    header = table.rows[0]
    repeat_header(header)
    prevent_row_split(header)
    for idx, (cell, label) in enumerate(zip(header.cells, headers)):
        set_cell_shading(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_run_font(r, font_size, WHITE, bold=True)
    for row_idx, values in enumerate(rows, start=1):
        row = table.add_row()
        prevent_row_split(row)
        for col_idx, (cell, value) in enumerate(zip(row.cells, values)):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if row_idx % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(str(value))
            set_run_font(
                r,
                font_size,
                BLUE if first_col_bold and col_idx == 0 else INK,
                bold=first_col_bold and col_idx == 0,
            )
            set_cell_border(cell, bottom={"val": "single", "sz": "3", "color": "D8DEE3"})


def add_status_strip(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    values = [
        ("63,3%", "ÍNDICE INFORMADO", BLUE),
        ("INTERMEDIÁRIO", "NÍVEL INFORMADO", NAVY),
        ("85,7%", "ESSENCIAIS INFORMADOS", RED),
        ("+21,4 p.p.", "EVOLUÇÃO INFORMADA", GREEN),
    ]
    set_table_geometry(table, [2409, 2409, 2409, 2411], indent_dxa=0)
    for cell, (value, label, fill) in zip(table.rows[0].cells, values):
        set_cell_shading(cell, fill)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        value_size = 11.5 if len(value) > 10 else 16
        set_run_font(r, value_size, WHITE, bold=True)
        r = p.add_run(f"\n{label}")
        set_run_font(r, 7.5, WHITE, bold=True)
    prevent_row_split(table.rows[0])


def load_criteria() -> list[dict[str, str]]:
    allowed = {
        "COMUM",
        "COMUM (EXCETO ESTATAIS INDEPENDENTES)",
        "COMUM (EXCETO ESTATAIS)",
        "PODER LEGISLATIVO",
    }
    with CSV_PATH.open(encoding="utf-8-sig", newline="") as stream:
        rows = [row for row in csv.DictReader(stream) if row["matriz"] in allowed]

    def natural_key(value: str) -> tuple[int, ...]:
        return tuple(int(number) for number in re.findall(r"\d+", value))

    rows.sort(key=lambda row: natural_key(row["id"]))
    return rows


def evidence_for(row: dict[str, str]) -> str:
    dimension = row["dimensao"]
    criterion = row["criterio"].lower()
    if dimension in {"Receita", "Despesa"}:
        return "URL direta; consulta por exercício; CSV/XLSX/JSON; data/hora de atualização; campos completos."
    if dimension == "Licitações":
        return "Processo integral; pesquisa e filtros; anexos; situação; datas; vínculo com contrato, PNCP e despesa."
    if dimension == "Contratos":
        return "Relação sequencial; inteiro teor; aditivos; fiscal; vigência; execução; pagamentos; vínculo PNCP."
    if dimension == "Obras":
        return "Painel da obra; contrato; medições; quantitativos; preços; percentual; fotos e situação atual."
    if dimension == "Recursos Humanos":
        return "Base nominal pesquisável; competência; histórico; exportação; ato de ingresso/saída e tabela legal."
    if dimension == "Diárias":
        return "Beneficiário, cargo, destino, motivo, período, quantidade, valor, norma e exportação."
    if dimension == "Serviço de Informação ao Cidadão - SIC":
        return "Página própria; formulário funcional; protocolo; prazos; recursos; relatórios e listas de sigilo."
    if dimension == "Acessibilidade":
        return "Teste em site e portal; teclado; leitor de tela; contraste; redimensionamento; mapa/caminho."
    if dimension == "Ouvidorias":
        return "Canal presencial e digital; protocolo; prazo; acompanhamento; Carta de Serviços atualizada."
    if dimension == "Lei Geral de Proteção de Dados (LGPD) e Governo Digital":
        return "Responsável e canal; política; serviços digitais; arquivos abertos/API; regras de uso; pesquisa."
    if dimension == "Atividades Finalísticas - PL":
        return "SAPL/site com URL permanente; documentos e anexos; autoria; tramitação; presença; votação e mídia."
    if dimension == "Convênios e Transferências":
        return "Relação, partes, objeto, vigência, valores, execução e inteiro teor; declaração negativa quando cabível."
    if dimension == "Planejamento e Prestação de contas":
        return "Documento integral; exercício/período; data de publicação; histórico; decisão do TCE e exportação."
    if dimension == "Informações Institucionais":
        return "Página atualizada, ato de suporte, responsável, data de revisão, contatos e histórico."
    if dimension == "Informações Prioritárias":
        return "Acesso visível na capa, busca funcional, endereço estável e teste sem cadastro."
    if "pesquisa" in criterion:
        return "Teste de busca textual e por filtros, com resultado exportável e sem erro."
    return "URL direta, documento integral, data de atualização, responsável e histórico."


def action_for(row: dict[str, str]) -> str:
    special = {
        "1.1": "Manter o sítio oficial estável, seguro, responsivo e claramente identificado como canal institucional.",
        "1.2": "Manter portal próprio ou compartilhado com acesso público, navegação estável e identificação inequívoca da Câmara.",
        "1.3": "Dar destaque permanente ao botão Transparência na capa e em todas as versões móvel e desktop.",
        "1.4": "Unificar busca do site e do portal, com pesquisa textual, filtros e orientação quando não houver resultado.",
        "3.1": "Disponibilizar previsão e realização das receitas por exercício, natureza, fonte e data, com exportação integral.",
        "4.1": "Exibir totais empenhados, liquidados e pagos, conciliados e atualizados até o primeiro dia útil após o registro.",
        "4.2": "Permitir consulta da despesa por órgão/unidade, função, programa, ação, natureza, fonte e exercício.",
        "4.3": "Ligar cada empenho ao credor, CNPJ/CPF protegido quando necessário, objeto, valor, licitação, contrato e pagamentos.",
        "8.3": "Publicar documentos das fases preparatória e externa, preservando estudo técnico, termo de referência, pareceres, propostas, atas e decisões.",
        "8.4": "Organizar dispensas e inexigibilidades por processo, com justificativas, pareceres, ratificação, contratação e PNCP.",
        "9.4": "Publicar ordem cronológica por fonte de recursos e justificar de forma individualizada qualquer alteração.",
        "10.4": "Criar relação específica de obras paralisadas; quando inexistentes, publicar declaração negativa datada.",
        "11.5": "Publicar RGF no prazo legal, com todos os anexos, assinatura, período e série histórica.",
        "12.4": "Simplificar o e-SIC, solicitando apenas dados necessários e eliminando exigência de documento, assinatura ou justificativa.",
        "12.7": "Publicar anualmente estatísticas de pedidos recebidos, atendidos, indeferidos, prazos médios e perfil genérico.",
        "12.8": "Publicar lista de informações classificadas; se não houver, divulgar declaração negativa anual.",
        "12.9": "Publicar lista das informações desclassificadas nos últimos 12 meses; se não houver, divulgar declaração negativa.",
        "14.3": "Publicar Carta de Serviços com serviços, requisitos, etapas, canais, prazos, compromissos e padrão de atendimento.",
        "15.4": "Oferecer CSV, XLSX ou JSON, metadados, dicionário e API ou endpoint documentado, com regras claras de reutilização.",
        "15.5": "Informar se houve adoção municipal da Lei nº 14.129/2021 e publicar a norma; se ausente, propor regulamentação local.",
        "20.3": "Garantir tramitação legislativa completa: ementa, autoria, relatoria, anexos, pareceres, emendas, situação e histórico.",
        "20.5": "Publicar pautas das comissões com antecedência e manter arquivo por comissão, data e matéria.",
        "20.6": "Publicar atas e presenças em formato pesquisável, ligadas à sessão, pauta, vídeo e votações.",
        "20.7": "Criar consulta de votações nominais por matéria, sessão, vereador e resultado.",
        "20.8": "Reunir decreto legislativo, parecer prévio do TCE, ata e resultado do julgamento das contas do Executivo.",
        "20.9": "Transmitir e arquivar sessões e audiências, com data, pauta, descrição, capítulos e link permanente.",
        "20.10": "Publicar norma, limites e valores de cotas/verbas, com prestação por parlamentar, competência e comprovantes.",
        "20.11": "Criar painel de atividades por parlamentar: proposições, relatorias, presenças, votações, comissões e despesas.",
    }
    criterion_id = row["id"].rstrip(".")
    if criterion_id in special:
        return special[criterion_id]
    text = row["criterio"].rstrip("?")
    replacements = (
        ("Possui ", "Manter e comprovar "),
        ("Existe ", "Manter e identificar "),
        ("Há ", "Disponibilizar e testar "),
        ("Divulga ", "Publicar e manter atualizada "),
        ("Publica ", "Publicar no prazo e preservar "),
        ("Participa ", "Manter presença institucional e divulgar "),
        ("Inclui ", "Incluir e manter "),
        ("Identifica ", "Identificar e manter atualizados "),
        ("Possibilita ", "Disponibilizar e testar "),
        ("Regulamenta ", "Regulamentar e divulgar "),
        ("Realiza ", "Realizar periodicamente e divulgar "),
        ("O acesso ", "Garantir que o acesso "),
        ("O site ", "Garantir que o site "),
        ("A solicitação ", "Garantir que a solicitação "),
        ("Contém ", "Disponibilizar "),
    )
    for old, new in replacements:
        if text.startswith(old):
            text = new + text[len(old):]
            break
    return text[0].upper() + text[1:] + "."


def add_main_content(doc: Document, criteria: list[dict[str, str]]) -> None:
    doc.add_heading("1. Síntese executiva", level=1)
    add_body(
        doc,
        "O resultado público anteriormente consultado para a Câmara indicou índice geral de "
        "63,3%, classificação “Intermediário”, atendimento de 85,7% dos critérios essenciais e "
        "evolução de 21,4 pontos percentuais. A melhora é relevante, mas o resultado confirma que "
        "o portal ainda não atende integralmente ao padrão esperado, especialmente porque os "
        "critérios essenciais não alcançaram 100%."
    )
    add_status_strip(doc)
    add_body(
        doc,
        "A página institucional de transparência já reúne links para contas, pessoal, licitações, "
        "contratos, planejamento, SIC, acessibilidade e LGPD. O principal desafio agora é transformar "
        "essa estrutura em uma experiência verificável: informação atualizada, completa, pesquisável, "
        "exportável, com histórico, data de publicação e ligação entre os sistemas."
    )

    counts = Counter(row["classe"] for row in criteria)
    add_callout(
        doc,
        "Escopo completo",
        f"A matriz PNTP 2026 contém 83 critérios aplicáveis à Câmara: {counts['Essencial']} "
        f"essenciais, {counts['Obrigatória']} obrigatórios e {counts['Recomendada']} recomendados. "
        "O Anexo A converte todos eles em ações e evidências de entrega, com prazo máximo de "
        "regularização até o 60º dia.",
        fill=PALE_GOLD,
        accent=GOLD,
    )
    spacer = doc.paragraphs[-1]
    spacer._element.getparent().remove(spacer._element)
    add_callout(
        doc,
        "Meta de 60 dias",
        "Até o encerramento do período, todos os 83 critérios deverão estar validados; todas as "
        "publicações atrasadas deverão estar regularizadas; e cada critério aplicável deverá ter "
        "evidência, responsável, data de atualização e teste final registrados.",
        fill=PALE_GREEN,
        accent=GREEN,
    )
    spacer = doc.paragraphs[-1]
    spacer._element.getparent().remove(spacer._element)

    doc.add_heading("2. Como os apontamentos foram classificados", level=1)
    add_table(
        doc,
        ["CLASSIFICAÇÃO", "SIGNIFICADO", "COMO TRATAR"],
        [
            [
                "Falha confirmada",
                "Resultado público ou ausência objetivamente verificável.",
                "Corrigir, publicar a evidência e registrar a data da regularização.",
            ],
            [
                "Deficiência observada",
                "Barreira encontrada na navegação, organização ou forma de entrega.",
                "Validar com o setor responsável e ajustar a experiência do usuário.",
            ],
            [
                "Atraso relatado",
                "Entrega ou publicação informada como pendente ou fora da rotina esperada.",
                "Inventariar o passivo, definir responsável e publicar cronograma de recuperação.",
            ],
            [
                "Item a validar",
                "Critério aplicável cujo espelho detalhado não foi disponibilizado.",
                "Conferir com a equipe e o TCE-MG, sem presumir falha antes da validação.",
            ],
        ],
        [2100, 3500, 4038],
        font_size=9.0,
        first_col_bold=True,
    )

    doc.add_heading("3. Apontamento das principais falhas e deficiências", level=1)
    findings = [
        [
            "F-01",
            "Falha confirmada",
            "Índice geral de 63,3%, ainda no nível Intermediário.",
            "Publicar plano critério a critério para alcançar pelo menos 75% e, depois, os níveis Prata/Ouro.",
            "Imediata",
        ],
        [
            "F-02",
            "Falha confirmada",
            "Critérios essenciais informados em 85,7%, abaixo dos 100% exigidos para selo.",
            "Identificar o essencial não atendido, corrigir primeiro e anexar evidência verificável.",
            "Imediata",
        ],
        [
            "F-03",
            "Atraso relatado",
            "Há entregas e publicações com atraso, sem inventário público consolidado do passivo.",
            "Criar força-tarefa, listar pendências por competência e publicar calendário de regularização.",
            "0-30 dias",
        ],
        [
            "F-04",
            "Deficiência observada",
            "Informações estão distribuídas entre site institucional, Betha, SAPL, PNCP e páginas de downloads.",
            "Criar página única por tema, com links diretos, contexto, responsável e data de atualização.",
            "0-30 dias",
        ],
        [
            "F-05",
            "Deficiência observada",
            "A página central apresenta links, mas não mostra status, cobertura histórica e última atualização de cada base.",
            "Acrescentar painel de atualização, período disponível, formato e unidade responsável.",
            "0-30 dias",
        ],
        [
            "F-06",
            "Deficiência observada",
            "O portal financeiro depende de aplicação JavaScript e não oferece, na página central, alternativa clara de dados abertos.",
            "Disponibilizar exportações integrais e endpoints documentados, além de páginas indexáveis e acessíveis.",
            "31-45 dias",
        ],
        [
            "F-07",
            "Item a validar",
            "Não há espelho público critério a critério com evidência, justificativa e situação da correção.",
            "Publicar matriz PNTP com status, URL, evidência, responsável, prazo e histórico de revisão.",
            "0-30 dias",
        ],
        [
            "F-08",
            "Item a validar",
            "Busca, filtros, exportação, série histórica e gravação de relatórios precisam ser testados em todas as bases.",
            "Executar roteiro funcional e corrigir cada componente que reduza a pontuação.",
            "31-45 dias",
        ],
        [
            "F-09",
            "Item a validar",
            "Dados de contratos, despesas, licitações e pagamentos podem existir sem identificador comum.",
            "Adotar chaves que permitam cruzar processo, contrato, empenho, liquidação, pagamento e PNCP.",
            "31-60 dias",
        ],
        [
            "F-10",
            "Deficiência de governança",
            "Não está visível uma rotina pública de responsáveis, prazos e testes periódicos do portal.",
            "Instituir comitê intersetorial, calendário, indicadores e revisão mensal das evidências.",
            "0-10 dias",
        ],
    ]
    add_table(
        doc,
        ["ID", "EVIDÊNCIA", "APONTAMENTO", "MELHORIA COLABORATIVA", "PRAZO"],
        findings,
        [620, 1500, 2580, 3938, 1000],
        font_size=8.2,
        first_col_bold=True,
    )
    add_callout(
        doc,
        "Cautela metodológica",
        "Como o espelho detalhado da avaliação não está disponível na página pública consultada, "
        "o documento não atribui automaticamente cada um dos 83 critérios como falha. A Câmara "
        "pode validar o anexo, indicar os itens já atendidos e concentrar a correção nos itens "
        "parciais, atrasados ou não atendidos.",
        fill=PALE_RED,
        accent=RED,
    )

    doc.add_heading("4. Plano de regularização integral em 60 dias", level=1)
    roadmap = [
        [
            "Dias 1-10",
            "Governança e diagnóstico",
            "Nomear ponto focal; reunir todos os setores; validar os 83 critérios; listar atrasos, links quebrados, documentos ausentes e dependências técnicas.",
            "100% dos critérios classificados, com responsável e prazo.",
        ],
        [
            "Dias 11-20",
            "Essenciais e passivo crítico",
            "Corrigir links; completar receitas, despesas, empenhos e RGF; publicar documentos essenciais e iniciar a recuperação das entregas vencidas.",
            "100% dos essenciais com URL e evidência testada.",
        ],
        [
            "Dias 21-30",
            "Publicações atrasadas",
            "Regularizar integralmente o passivo de publicações; completar pessoal, diárias, licitações, contratos, obras, prestação de contas e documentos administrativos.",
            "100% das publicações atrasadas disponíveis e datadas.",
        ],
        [
            "Dias 31-40",
            "Dados abertos e serviços",
            "Implantar busca, filtros, exportação, metadados, histórico, SIC, Ouvidoria, Carta de Serviços, LGPD e acessibilidade.",
            "Bases pesquisáveis, exportáveis e acessíveis.",
        ],
        [
            "Dias 41-50",
            "Integração legislativa",
            "Ligar SAPL, Betha e PNCP; completar pautas, comissões, atas, presenças, votações, julgamento de contas, transmissões e dados parlamentares.",
            "Atividade legislativa completa e rastreável.",
        ],
        [
            "Dias 51-60",
            "Teste, correção e entrega",
            "Testar os 83 critérios, corrigir falhas remanescentes, publicar matriz final, evidências, histórico e relatório de conclusão; realizar validação cidadã.",
            "Plano concluído, sem pendência aplicável aberta.",
        ],
        [
            "Após o dia 60",
            "Manutenção preventiva",
            "Manter revisão mensal, alertas de prazo, teste trimestral e registro público de atualizações e regressões.",
            "Impedir novos atrasos sem postergar a conclusão do plano.",
        ],
    ]
    add_table(
        doc,
        ["JANELA", "FOCO", "ENTREGAS", "RESULTADO ESPERADO"],
        roadmap,
        [1250, 1800, 4438, 2150],
        font_size=8.8,
        first_col_bold=True,
    )
    add_callout(
        doc,
        "Regra de conclusão",
        "Uma ação somente será considerada concluída quando houver publicação acessível, URL direta, "
        "conteúdo integral, data de atualização, responsável identificado e teste de funcionamento. "
        "Dependência de fornecedor deverá ter solução alternativa e não poderá manter a informação "
        "pública indisponível após o 60º dia.",
        fill=PALE_GREEN,
        accent=GREEN,
    )

    doc.add_heading("5. Pacotes de melhoria por área", level=1)
    packages = [
        ["Governança do portal", "Matriz pública, responsáveis, calendário, rotina de atualização, registro de correções e teste mensal."],
        ["Arquitetura e navegação", "Página única por tema, URLs diretas e permanentes, busca, mapa, caminho de páginas e aviso de indisponibilidade."],
        ["Atualidade e prazos", "Data/hora de registro e publicação, frequência de atualização, passivo de atrasos e alertas antes do vencimento."],
        ["Dados abertos", "CSV/XLSX/JSON, dicionário, metadados, licença, API, exportação integral e série histórica."],
        ["Receita e despesa", "Previsão/realização; empenho/liquidação/pagamento; credor, objeto, classificação, processo, licitação e contrato."],
        ["Licitações e contratos", "Processo integral, contratação direta, PNCP, aditivos, fiscais, execução, sanções e ordem cronológica."],
        ["Obras", "Situação, prazos, empresa, quantitativos, preços, medições, pagamentos, fotos e paralisações."],
        ["Pessoal, diárias e verbas", "Base nominal, remuneração, lotação, atos, históricos, diárias detalhadas e cotas parlamentares."],
        ["Planejamento e contas", "Balanço, RGF, decisões do TCE, plano estratégico, relatório de atividades e histórico."],
        ["SIC e Ouvidoria", "Canais simples, protocolos, prazos, recursos, relatórios estatísticos, Carta de Serviços e acompanhamento."],
        ["Acessibilidade", "Teclado, leitor de tela, contraste, redimensionamento, linguagem simples, responsividade e documentos acessíveis."],
        ["LGPD e Governo Digital", "Encarregado, política, base legal, serviços digitais, adoção normativa e proteção sem ocultação genérica."],
        ["Atividade legislativa", "Biografias, matérias, tramitação, pautas, comissões, atas, presenças, votações, contas e transmissões."],
        ["Participação cidadã", "Consultas, audiências, pesquisa de satisfação, devolutiva e publicação das contribuições recebidas."],
    ]
    add_table(
        doc,
        ["ÁREA", "MELHORIAS PROPOSTAS"],
        packages,
        [2550, 7088],
        font_size=9.0,
        first_col_bold=True,
    )

    doc.add_heading("6. Padrão mínimo para cada publicação", level=1)
    standards = [
        ["Identificação", "Título claro, órgão/unidade, responsável, exercício ou competência e descrição do conteúdo."],
        ["Atualidade", "Data/hora do registro, data/hora da publicação, última atualização e periodicidade esperada."],
        ["Completude", "Documento integral e anexos; declaração negativa quando o evento não ocorreu."],
        ["Formato", "Consulta em tela e exportação integral em CSV, XLSX ou JSON; PDF apenas como apresentação complementar."],
        ["Pesquisa", "Busca textual, filtros combináveis, ordenação e paginação transparente."],
        ["Histórico", "Série mínima de três exercícios ou o período exigido; versões anteriores preservadas."],
        ["Integração", "Identificadores que liguem matéria, processo, licitação, contrato, empenho, pagamento e PNCP."],
        ["Acessibilidade", "Navegação por teclado, leitor de tela, contraste, texto redimensionável e documento pesquisável."],
        ["Rastreabilidade", "URL permanente, versão, fonte, dicionário de dados e contato técnico."],
        ["Proteção de dados", "Tarja apenas no trecho protegido, entrega da parte pública e indicação da base legal da restrição."],
    ]
    add_table(
        doc,
        ["ELEMENTO", "PADRÃO DE ENTREGA"],
        standards,
        [2100, 7538],
        font_size=9.0,
        first_col_bold=True,
    )

    doc.add_heading("7. Prazos e periodicidades prioritários", level=1)
    deadlines = [
        ["Execução orçamentária e financeira", "Até o primeiro dia útil seguinte ao registro contábil.", "Decreto nº 10.540/2020."],
        ["Pedido de acesso à informação", "Imediato; se inviável, até 20 dias, prorrogáveis uma vez por 10 dias com justificativa.", "LAI e Lei Municipal nº 5.729/2013."],
        ["Recurso LAI", "Requerente: 10 dias; decisão da autoridade superior: 5 dias.", "Lei nº 12.527/2011, art. 15."],
        ["Ouvidoria", "Decisão final em 30 dias, prorrogável uma vez por mais 30, com justificativa.", "Lei nº 13.460/2017, art. 16."],
        ["RGF", "Até 30 dias após cada quadrimestre.", "LRF, arts. 54 e 55."],
        ["Contas municipais", "Durante todo o exercício; também por 60 dias anuais para exame do contribuinte.", "LRF, art. 49; CF, art. 31, §3º."],
        ["Contratos no PNCP", "20 dias úteis após assinatura em licitação; 10 dias úteis na contratação direta.", "Lei nº 14.133/2021, art. 94."],
        ["Obras contratadas", "Quantidades/preços em 25 dias úteis da assinatura; executados/pagos em 45 dias úteis da conclusão.", "Lei nº 14.133/2021, art. 94, §3º."],
    ]
    add_table(
        doc,
        ["INFORMAÇÃO", "PRAZO", "BASE"],
        deadlines,
        [2600, 4388, 2650],
        font_size=8.7,
        first_col_bold=True,
    )

    doc.add_heading("8. Modelo de governança participativa", level=1)
    governance = [
        ["Mesa Diretora", "Patrocinar o plano, remover impedimentos e acompanhar o cronograma mensal."],
        ["Controladoria Interna", "Validar critérios, evidências, prazos legais e correções antes da publicação."],
        ["Tecnologia da Informação", "Garantir disponibilidade, integrações, acessibilidade, busca, exportação, APIs e monitoramento."],
        ["Contabilidade e Finanças", "Manter receitas, despesas, RGF, balanço, pagamentos e conciliações no prazo."],
        ["Licitações e Contratos", "Publicar processos, documentos, PNCP, contratos, aditivos, fiscais e execução."],
        ["Recursos Humanos", "Atualizar pessoal, remuneração, cargos, atos, concursos, estagiários e terceirizados."],
        ["Secretaria Legislativa", "Atualizar matérias, pautas, comissões, atas, presença, votações e julgamento de contas."],
        ["SIC, Ouvidoria e LGPD", "Cuidar dos canais, prazos, relatórios, Carta de Serviços, privacidade e acesso parcial."],
        ["Comunicação", "Traduzir conteúdos, divulgar atualizações e ampliar participação e pesquisa de satisfação."],
        ["Sociedade civil", "Testar o portal, indicar barreiras, participar de reuniões e acompanhar a matriz pública."],
    ]
    add_table(
        doc,
        ["RESPONSÁVEL", "CONTRIBUIÇÃO SUGERIDA"],
        governance,
        [2500, 7138],
        font_size=9.0,
        first_col_bold=True,
    )
    add_callout(
        doc,
        "Proposta de participação",
        "Realizar uma reunião técnica inicial, seguida de devolutiva pública em até 30 dias. "
        "O Fiscaliza Varginha pode contribuir com testes de navegação, organização das evidências "
        "e comunicação em linguagem cidadã, respeitando a autonomia administrativa da Câmara.",
        fill=PALE_GREEN,
        accent=GREEN,
    )

    doc.add_heading("9. Painel público de acompanhamento", level=1)
    add_body(
        doc,
        "Sugere-se publicar uma planilha ou painel com uma linha para cada critério. O acompanhamento "
        "deve permanecer acessível entre os ciclos de avaliação e registrar tanto as correções quanto "
        "eventuais regressões."
    )
    tracker = [
        ["Critério PNTP", "ID, classe e descrição."],
        ["Situação", "Atende; atende parcialmente; não atende; não se aplica; em validação."],
        ["Evidência", "URL direta, arquivo, tela ou relatório que comprova o atendimento."],
        ["Qualidade", "Disponibilidade, atualidade, histórico, gravação/exportação e pesquisa."],
        ["Responsável", "Setor e agente encarregado pela atualização."],
        ["Prazo", "Data prevista e data efetiva de conclusão."],
        ["Observação", "Pendência, justificativa, dependência do fornecedor e próxima ação."],
        ["Revisão", "Data, responsável e resultado do último teste."],
    ]
    add_table(
        doc,
        ["CAMPO", "CONTEÚDO"],
        tracker,
        [2200, 7438],
        font_size=9.0,
        first_col_bold=True,
    )

    doc.add_heading("10. Base normativa e referências", level=1)
    references = [
        ("Portal de Transparência da Câmara", PORTAL_URL, "estrutura atual de acesso às informações"),
        ("Downloads oficiais PNTP", PNTP_DOWNLOADS_URL, "matriz do ciclo 2026"),
        ("Cartilha PNTP 2025 - Atricon", PNTP_CARTILHA_URL, "metodologia e níveis de transparência"),
        ("Resultado público anteriormente consultado", RESULT_URL, "índice e evolução; sujeito à confirmação oficial"),
        ("Lei nº 12.527/2011 - LAI", LAI_URL, "transparência ativa, formatos, prazos e recursos"),
        ("Lei Municipal nº 5.729/2013", LOCAL_LAI_URL, "regulamentação local para Executivo e Legislativo"),
        ("Lei Complementar nº 101/2000 - LRF", LRF_URL, "contas, relatórios e fiscalização"),
        ("Decreto nº 10.540/2020 - SIAFIC", SIAFIC_URL, "tempo real e detalhamento"),
        ("Lei nº 14.133/2021", PROCUREMENT_URL, "PNCP, contratos e obras"),
        ("Lei nº 13.460/2017", OUVIDORIA_URL, "Ouvidoria e direitos do usuário"),
        ("Lei nº 13.709/2018 - LGPD", LGPD_URL, "proteção de dados pelo poder público"),
        ("Lei nº 13.146/2015 - LBI", LBI_URL, "acessibilidade dos sítios governamentais"),
    ]
    for idx, (label, url, note) in enumerate(references, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.15)
        p.paragraph_format.first_line_indent = Cm(-0.15)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{idx}. ")
        set_run_font(r, 9.8, BLUE, bold=True)
        add_hyperlink(p, label, url)
        r = p.add_run(f" - {note}.")
        set_run_font(r, 9.8, INK)


def add_annex(doc: Document, criteria: list[dict[str, str]]) -> None:
    doc.add_heading("Anexo A - Checklist completo PNTP 2026 para a Câmara", level=1)
    add_body(
        doc,
        "Este anexo reúne todos os 83 critérios aplicáveis às matrizes COMUM, COMUM (exceto "
        "estatais), COMUM (exceto estatais independentes) e PODER LEGISLATIVO. A coluna Situação "
        "sugerida não presume falha: orienta a Câmara a validar, regularizar e publicar a evidência."
    )
    add_callout(
        doc,
        "Regra de priorização",
        "Critérios essenciais devem alcançar 100%. Em seguida, corrigem-se os obrigatórios e, "
        "depois, os recomendados. Cada validação deve testar disponibilidade, atualidade, série "
        "histórica, gravação/exportação e filtros de pesquisa.",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    grouped: OrderedDict[str, list[dict[str, str]]] = OrderedDict()
    for row in criteria:
        grouped.setdefault(row["dimensao"], []).append(row)

    for dimension, rows in grouped.items():
        heading = doc.add_heading(dimension, level=2)
        heading.paragraph_format.keep_with_next = True
        table = doc.add_table(rows=1, cols=5)
        widths = [620, 1250, 3100, 2868, 1800]
        set_table_geometry(table, widths)
        repeat_header(table.rows[0])
        prevent_row_split(table.rows[0])
        headers = ["ID", "CLASSE", "CRITÉRIO PNTP", "AÇÃO DE MELHORIA", "EVIDÊNCIA ESPERADA"]
        for cell, label in zip(table.rows[0].cells, headers):
            set_cell_shading(cell, NAVY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(label)
            set_run_font(r, 7.5, WHITE, bold=True)
        for idx, row_data in enumerate(rows, start=1):
            row = table.add_row()
            prevent_row_split(row)
            values = [
                row_data["id"],
                row_data["classe"],
                row_data["criterio"],
                action_for(row_data),
                evidence_for(row_data),
            ]
            for col_idx, (cell, value) in enumerate(zip(row.cells, values)):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                if idx % 2 == 0:
                    set_cell_shading(cell, LIGHT_GRAY)
                if col_idx == 1:
                    fill = (
                        PALE_RED if row_data["classe"] == "Essencial"
                        else PALE_BLUE if row_data["classe"] == "Obrigatória"
                        else PALE_GREEN
                    )
                    set_cell_shading(cell, fill)
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                r = p.add_run(value)
                color = (
                    RED if row_data["classe"] == "Essencial"
                    else BLUE if row_data["classe"] == "Obrigatória"
                    else GREEN
                )
                set_run_font(
                    r,
                    8.2,
                    color if col_idx in (0, 1) else INK,
                    bold=col_idx in (0, 1),
                )
                set_cell_border(cell, bottom={"val": "single", "sz": "3", "color": "D8DEE3"})
        note = doc.add_paragraph()
        note.paragraph_format.space_before = Pt(4)
        note.paragraph_format.space_after = Pt(4)
        r = note.add_run(
            "Situação inicial sugerida: VALIDAR/REGULARIZAR - concluir até o 60º dia e registrar URL, data, responsável e evidência."
        )
        set_run_font(r, 8.5, MID_GRAY, italic=True)


def add_closing(doc: Document) -> None:
    doc.add_heading("Encaminhamento colaborativo", level=1)
    add_body(
        doc,
        "Sugere-se que este plano seja recebido como contribuição técnica e cidadã, submetido aos "
        "setores responsáveis e executado em até 60 dias, com indicação dos itens já atendidos, das "
        "correções aceitas, dos responsáveis e das datas de conclusão. A publicação da própria "
        "resposta e do cronograma reforçará a confiança, facilitará o controle interno e permitirá "
        "que a sociedade acompanhe os avanços de forma construtiva."
    )
    add_callout(
        doc,
        "Resultado pretendido",
        "Um portal confiável, atualizado, acessível e reutilizável, capaz de apoiar o cidadão, a "
        "gestão, os vereadores, o controle interno e os órgãos de fiscalização - com menos retrabalho "
        "e maior clareza sobre cada entrega.",
        fill=PALE_GREEN,
        accent=GREEN,
    )
    # The callout helper adds a neutral spacer for normal section flow. Remove it
    # here so the closing signature stays with the final message.
    spacer = doc.paragraphs[-1]
    spacer._element.getparent().remove(spacer._element)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.keep_together = True
    r = p.add_run("FISCALIZA VARGINHA")
    set_run_font(r, 11, NAVY, bold=True)
    r = p.add_run("\nIniciativa cidadã independente e apartidária")
    set_run_font(r, 9.5, MID_GRAY)


def build() -> Path:
    criteria = load_criteria()
    if len(criteria) != 83:
        raise ValueError(f"Esperados 83 critérios aplicáveis, encontrados {len(criteria)}")
    doc = Document()
    configure_document(doc)
    add_title_block(doc)
    add_main_content(doc, criteria)
    add_annex(doc, criteria)
    add_closing(doc)

    core = doc.core_properties
    core.title = "Plano Colaborativo de Melhoria do Portal da Transparência da Câmara de Varginha"
    core.subject = "Diagnóstico e regularização integral do portal em 60 dias"
    core.author = "Fiscaliza Varginha"
    core.keywords = "transparência, Câmara Municipal, Varginha, PNTP, controle social"
    core.comments = "Documento colaborativo, sem presunção de irregularidade."
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build())
