from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "pntp" / "solicitacoes-portal"
OUT_DIR.mkdir(parents=True, exist_ok=True)

NAVY = "17324D"
BLUE = "1F5D8F"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "687581"
GREEN = "297A5B"
WHITE = "FFFFFF"
BLACK = "20262C"
RED = "A33A32"

LAI_URL = "https://www.planalto.gov.br/ccivil_03/_ato2011-2014/2011/lei/l12527.htm"
LRF_URL = "https://www.planalto.gov.br/ccivil_03/leis/lcp/lcp101.htm"
CONSTITUTION_URL = "https://www.planalto.gov.br/ccivil_03/constituicao/constituicaocompilado.htm"
LOCAL_LAI_URL = (
    "https://sapl.varginha.mg.leg.br/media/sapl/public/normajuridica/"
    "2013/5880/lei5729-2013.pdf"
)
SIAFIC_URL = "https://www.planalto.gov.br/ccivil_03/_ato2019-2022/2020/decreto/d10540.htm"
GOV_DIGITAL_URL = "https://www.planalto.gov.br/ccivil_03/_ato2019-2022/2021/lei/l14129.htm"
USER_SERVICE_URL = "https://www.planalto.gov.br/ccivil_03/_ato2015-2018/2017/lei/l13460.htm"
PROCUREMENT_URL = "https://www.planalto.gov.br/ccivil_03/_ato2019-2022/2021/lei/l14133.htm"
LGPD_URL = "https://www.planalto.gov.br/ccivil_03/_ato2015-2018/2018/lei/l13709compilado.htm"
ACCESSIBILITY_URL = "https://www.planalto.gov.br/ccivil_03/_ato2015-2018/2015/lei/l13146.htm"
ARCHIVES_URL = "https://www.planalto.gov.br/ccivil_03/leis/l8159.htm"
PNTP_URL = "https://radardatransparencia.atricon.org.br/pdf/Cartilha-PNTP-2025.pdf"
PNTP_DOWNLOADS_URL = "https://radardatransparencia.atricon.org.br/downloads.html"
PREF_RESULT_URL = "https://radartransparente.com.br/avaliacao/59f001b0-a3d0-42cb-953b-0d02db067387"
CAM_RESULT_URL = "https://radartransparente.com.br/avaliacao/5518baf1-b161-4f79-b364-94c64e2af7bb"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **edges) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name, attrs in edges.items():
        edge = tc_borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            tc_borders.append(edge)
        for key, value in attrs.items():
            edge.set(qn(f"w:{key}"), str(value))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_keep_with_next(paragraph, value: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def set_keep_lines(paragraph, value: bool = True) -> None:
    paragraph.paragraph_format.keep_together = value


def add_hyperlink(paragraph, text: str, url: str, color: str = BLUE) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    new_run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    run_color = OxmlElement("w:color")
    run_color.set(qn("w:val"), color)
    run_properties.append(run_color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)
    new_run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_bookmark(paragraph, name: str, bookmark_id: str) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bookmark_id)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bookmark_id)
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def configure_document(doc: Document, recipient_short: str) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.75)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.55)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.7)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(4.5)
    normal.paragraph_format.line_spacing = 1.04

    for style_name, size, color in [
        ("Title", 22, NAVY),
        ("Subtitle", 11, MID_GRAY),
        ("Heading 1", 14, NAVY),
        ("Heading 2", 11.5, BLUE),
        ("Heading 3", 10.5, NAVY),
    ]:
        style = styles[style_name]
        style.font.name = "Aptos Display" if style_name in {"Title", "Heading 1"} else "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Cm(16.9))
    table.autofit = False
    table.columns[0].width = Cm(10.7)
    table.columns[1].width = Cm(6.2)
    table.cell(0, 0).width = Cm(10.7)
    table.cell(0, 1).width = Cm(6.2)
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    set_cell_shading(left, NAVY)
    set_cell_shading(right, BLUE)
    for cell in (left, right):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.margin_top = Cm(0.12)
        cell.margin_bottom = Cm(0.12)
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("FISCALIZA VARGINHA")
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(WHITE)
    p2 = right.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(recipient_short.upper())
    r2.bold = True
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = RGBColor.from_string(WHITE)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(
        "Iniciativa cidadã independente • Transparência, dados públicos e controle social  |  "
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MID_GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run_page = p.add_run()
    run_page._r.append(fld_char1)
    run_page._r.append(instr_text)
    run_page._r.append(fld_char2)


def add_label_value(paragraph, label: str, value: str) -> None:
    run = paragraph.add_run(label)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)
    paragraph.add_run(value)


def add_cover(
    doc: Document,
    destination_lines: Iterable[str],
    title: str,
    subject: str,
    reference: str,
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("REQUERIMENTO ADMINISTRATIVO")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    r.font.letter_spacing = Pt(0.5)

    title_p = doc.add_paragraph(style="Title")
    title_p.paragraph_format.space_before = Pt(2)
    title_p.paragraph_format.space_after = Pt(4)
    title_p.add_run(title)

    sub = doc.add_paragraph(style="Subtitle")
    sub.paragraph_format.space_after = Pt(10)
    sub.add_run("Pedido de acesso à informação + solicitação colaborativa de providências")

    destination = doc.add_table(rows=1, cols=1)
    destination.autofit = False
    destination.columns[0].width = Cm(16.8)
    cell = destination.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_border(
        cell,
        left={"val": "single", "sz": "18", "color": BLUE},
        top={"val": "single", "sz": "4", "color": "C9D9E6"},
        bottom={"val": "single", "sz": "4", "color": "C9D9E6"},
        right={"val": "single", "sz": "4", "color": "C9D9E6"},
    )
    for idx, line in enumerate(destination_lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(line)
        r.bold = idx == 0
        r.font.size = Pt(9.5 if idx == 0 else 9)
        r.font.color.rgb = RGBColor.from_string(NAVY if idx == 0 else BLACK)

    meta = doc.add_table(rows=3, cols=2)
    meta.autofit = False
    meta.columns[0].width = Cm(3.1)
    meta.columns[1].width = Cm(13.7)
    meta_data = [
        ("ASSUNTO", subject),
        ("REFERÊNCIA", reference),
        ("DATA", "Varginha/MG, 25 de julho de 2026"),
    ]
    for row, (label, value) in zip(meta.rows, meta_data):
        row.cells[0].width = Cm(3.1)
        row.cells[1].width = Cm(13.7)
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for c in row.cells:
            set_cell_border(
                c,
                bottom={"val": "single", "sz": "4", "color": "D8DEE3"},
            )
        p0 = row.cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        rr = p0.add_run(label)
        rr.bold = True
        rr.font.size = Pt(8)
        rr.font.color.rgb = RGBColor.from_string(MID_GRAY)
        p1 = row.cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        p1.add_run(value)

    doc.add_paragraph()
    intro = doc.add_paragraph()
    add_label_value(intro, "Requerente: ", "[NOME COMPLETO / INSTITUIÇÃO]")
    intro.paragraph_format.space_after = Pt(1)
    intro2 = doc.add_paragraph()
    add_label_value(intro2, "Contato para resposta: ", "[E-MAIL]  •  [TELEFONE, se desejar]")
    intro2.paragraph_format.space_after = Pt(1)
    intro3 = doc.add_paragraph()
    add_label_value(intro3, "Número do protocolo: ", "[PREENCHER APÓS O ENVIO]")
    intro3.paragraph_format.space_after = Pt(8)


def add_summary_cards(doc: Document, cards: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=len(cards))
    table.autofit = False
    total = 16.8
    for col in table.columns:
        col.width = Cm(total / len(cards))
    for idx, (number, label, color) in enumerate(cards):
        cell = table.cell(0, idx)
        cell.width = Cm(total / len(cards))
        set_cell_shading(cell, color)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_border(
            cell,
            top={"val": "single", "sz": "4", "color": WHITE},
            bottom={"val": "single", "sz": "4", "color": WHITE},
            left={"val": "single", "sz": "4", "color": WHITE},
            right={"val": "single", "sz": "4", "color": WHITE},
        )
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(number)
        r.bold = True
        if len(cards) >= 4:
            number_size = 10.5 if len(number) > 10 else 13
        else:
            number_size = 16
        r.font.size = Pt(number_size)
        r.font.color.rgb = RGBColor.from_string(WHITE)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(label)
        r2.bold = True
        r2.font.size = Pt(7.8)
        r2.font.color.rgb = RGBColor.from_string(WHITE)


def add_body_text(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_keep_lines(p)
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix) :])
    else:
        p.add_run(text)


def add_numbered_items(doc: Document, items: list[tuple[str, str]], start: int = 1) -> None:
    for index, (title, body) in enumerate(items, start=start):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.15)
        p.paragraph_format.first_line_indent = Cm(-0.15)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_keep_lines(p)
        number = p.add_run(f"{index}. ")
        number.bold = True
        number.font.color.rgb = RGBColor.from_string(BLUE)
        heading = p.add_run(title)
        heading.bold = True
        p.add_run(body)


def add_priority_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    widths = [Cm(2.2), Cm(7.4), Cm(7.2)]
    headers = ["PRIORIDADE", "ENTREGA SUGERIDA", "RESULTADO ESPERADO"]
    for idx, (width, label) in enumerate(zip(widths, headers)):
        table.columns[idx].width = width
        cell = table.cell(0, idx)
        cell.width = width
        set_cell_shading(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(WHITE)
    set_repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for row_idx, values in enumerate(rows, start=1):
        row = table.add_row()
        prevent_row_split(row)
        cells = row.cells
        for idx, (cell, width, value) in enumerate(zip(cells, widths, values)):
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            r.font.size = Pt(8.5)
            if idx == 0:
                r.bold = True
                r.font.color.rgb = RGBColor.from_string(BLUE)
            set_cell_border(
                cell,
                bottom={"val": "single", "sz": "3", "color": "D8DEE3"},
            )


def add_legal_table(
    doc: Document,
    headers: tuple[str, str, str],
    rows: list[tuple[str, str, str]],
    widths: tuple[float, float, float] = (3.4, 7.0, 6.4),
) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    cm_widths = [Cm(value) for value in widths]
    for idx, (width, label) in enumerate(zip(cm_widths, headers)):
        table.columns[idx].width = width
        cell = table.cell(0, idx)
        cell.width = width
        set_cell_shading(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(7.7)
        r.font.color.rgb = RGBColor.from_string(WHITE)
    set_repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for row_idx, values in enumerate(rows, start=1):
        row = table.add_row()
        prevent_row_split(row)
        cells = row.cells
        for idx, (cell, width, value) in enumerate(zip(cells, cm_widths, values)):
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if row_idx % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(value)
            r.font.size = Pt(7.7)
            if idx == 0:
                r.bold = True
                r.font.color.rgb = RGBColor.from_string(BLUE)
            set_cell_border(
                cell,
                bottom={"val": "single", "sz": "3", "color": "D8DEE3"},
            )


def add_legal_framework(doc: Document, is_prefeitura: bool) -> None:
    doc.add_heading("6. Base legal detalhada para transparência e fiscalização", level=1)
    add_body_text(
        doc,
        "As normas abaixo estabelecem requisitos mínimos. A publicação não se considera "
        "plenamente atendida apenas porque existe um arquivo: a informação deve ser localizável, "
        "íntegra, atualizada, acessível, reutilizável e entregue no prazo aplicável."
    )

    add_legal_table(
        doc,
        ("NORMA", "DIREITO OU OBRIGAÇÃO", "REFLEXO PRÁTICO NO PORTAL"),
        [
            (
                "Constituição Federal\narts. 5º, XXXIII e XXXIV; 31; 37; 70 e 74, §2º",
                "Direito à informação e de petição; publicidade e eficiência; fiscalização "
                "municipal pelo Legislativo e controles internos; exame das contas por "
                "contribuintes; legitimidade cidadã para provocar o controle externo.",
                "Acesso sem barreiras indevidas, documentação apta à conferência e canais "
                "claros de pedido, recurso, representação e acompanhamento.",
            ),
            (
                "Lei Federal nº 12.527/2011 - LAI\narts. 3º, 7º a 12 e 15",
                "Publicidade é regra e sigilo é exceção. Garante informação primária, íntegra, "
                "autêntica e atualizada; acesso à parte não sigilosa; pesquisa, formatos abertos, "
                "acesso automatizado, integridade, atualização e acessibilidade.",
                "Entregar dados digitais no formato em que estão armazenados; oferecer CSV/XLSX/"
                "JSON quando houver base estruturada; não limitar a resposta a imagem ou PDF "
                "quando existirem dados reutilizáveis; ocultar somente a parte protegida.",
            ),
            (
                "Lei Municipal nº 5.729/2013\narts. 1º, 3º, 7º, 9º a 14 e 21 a 22",
                "Regulamenta a LAI em Varginha para Executivo e Legislativo; exige SIC, protocolo, "
                "transparência ativa, ferramenta de pesquisa, gravação em diversos formatos, "
                "informação sobre formatos e acessibilidade.",
                "O pedido pode ser físico ou virtual; deve gerar protocolo; a recusa precisa de "
                "fundamento legal e orientação de recurso. Aplica-se aos dois Poderes municipais.",
            ),
            (
                "LC nº 101/2000 - LRF\narts. 9º, §4º; 48, 48-A, 49, 52, 54, 55 e 59",
                "Ampla divulgação de PPA, LDO, LOA, contas, RREO e RGF; execução orçamentária "
                "pormenorizada; audiências de metas fiscais; fiscalização pelo Legislativo, "
                "Tribunal de Contas e controles internos.",
                "Dados devem permitir acompanhar receita, empenho, liquidação, pagamento, "
                "favorecido, objeto, processo, licitação, metas, limites e prestações de contas.",
            ),
            (
                "Decreto Federal nº 10.540/2020\narts. 2º, IX e X; 7º e 8º",
                "Define tempo real como publicação até o primeiro dia útil seguinte ao registro "
                "contábil no SIAFIC e exige meio eletrônico de amplo acesso público.",
                "Exibir data e hora do registro e da atualização; permitir consulta sem cadastro, "
                "senha ou autenticação; detalhar empenho, liquidação, pagamento e receitas.",
            ),
            (
                "Lei Federal nº 14.129/2021\narts. 2º, 4º e 29",
                "A aplicação direta aos Municípios depende de adoção por ato normativo próprio. "
                "Quando aplicável, define dados abertos estruturados, processáveis por máquina, "
                "referenciados na internet e disponíveis em formato aberto.",
                "Informar se houve adoção municipal. Independentemente disso, cumprir os formatos "
                "abertos já exigidos pela LAI: arquivo integral, dicionário, metadados, URL "
                "permanente e, quando possível, API documentada.",
            ),
            (
                "Lei Federal nº 14.133/2021\narts. 54, 94 e 174",
                "Edital e anexos no PNCP; documentos preparatórios após homologação; contratos e "
                "aditivos no PNCP como condição de eficácia, além de dados específicos de obras.",
                "Vincular no portal municipal edital, processo, resultado, contrato, aditivo, "
                "fiscal, execução, medições, pagamentos e registro correspondente no PNCP.",
            ),
            (
                "Leis Federais nº 13.709/2018 e nº 13.146/2015\nLGPD art. 23; LBI art. 63",
                "O poder público deve informar finalidade, base legal e práticas de tratamento de "
                "dados pessoais; os sítios governamentais devem ser acessíveis.",
                "LGPD não autoriza ocultação genérica: proteger dados pessoais e fornecer a parte "
                "pública. Garantir teclado, leitor de tela, contraste, texto redimensionável e "
                "símbolo de acessibilidade.",
            ),
            (
                "Lei Federal nº 8.159/1991\narts. 1º a 5º e 7º",
                "Impõe gestão e proteção dos documentos públicos como elementos de prova e "
                "informação e reconhece arquivos municipais dos Poderes Executivo e Legislativo.",
                "Preservar histórico, anexos, versões, links permanentes e trilha de atualização; "
                "evitar substituição de arquivos sem registro e remoção indevida de séries antigas.",
            ),
            (
                "Lei Federal nº 13.460/2017\narts. 13 a 16",
                "Disciplina manifestações de Ouvidoria e determina decisão final ao usuário.",
                "Separar claramente pedido LAI de manifestação de Ouvidoria e publicar protocolo, "
                "canal, andamento e prazo próprio de cada procedimento.",
            ),
        ],
    )

    doc.add_heading("7. Prazos objetivos que devem ser demonstrados", level=1)
    deadline_rows = [
        (
            "Pedido LAI",
            "Imediato, se disponível; caso contrário, até 20 dias, prorrogáveis uma vez por 10 dias, com justificativa expressa.",
            "LAI, art. 11; Lei Municipal nº 5.729/2013, art. 13.",
        ),
        (
            "Recurso LAI",
            "Interposição pelo requerente em até 10 dias; decisão da autoridade superior em 5 dias, conforme a LAI federal.",
            "LAI, art. 15. A lei municipal também prevê recurso em 10 dias e reclamação subsequente em 5 dias.",
        ),
        (
            "Manifestação de Ouvidoria",
            "Decisão final em 30 dias, prorrogável uma única vez por igual período, mediante justificativa.",
            "Lei nº 13.460/2017, art. 16. Este prazo não substitui o prazo da LAI.",
        ),
        (
            "Execução orçamentária e financeira",
            "Até o primeiro dia útil seguinte à data do registro contábil no SIAFIC.",
            "Decreto nº 10.540/2020, arts. 2º, IX, e 7º, §1º.",
        ),
        (
            "RREO",
            "Até 30 dias após o encerramento de cada bimestre.",
            "LRF, art. 52.",
        ),
        (
            "RGF",
            "Até 30 dias após o encerramento de cada quadrimestre, com amplo acesso inclusive eletrônico.",
            "LRF, arts. 54 e 55, §2º.",
        ),
        (
            "Metas fiscais",
            "Audiências até o final de fevereiro, maio e setembro, referentes ao quadrimestre anterior.",
            "LRF, art. 9º, §4º.",
        ),
        (
            "Programação financeira",
            "Estabelecimento da programação financeira e do cronograma mensal de desembolso em até 30 dias após a publicação do orçamento.",
            "LRF, art. 8º.",
        ),
        (
            "Contas municipais",
            "Disponíveis durante todo o exercício; adicionalmente, por 60 dias anuais para exame de qualquer contribuinte.",
            "LRF, art. 49; Constituição Federal, art. 31, §3º.",
        ),
        (
            "Contrato e aditamento no PNCP",
            "Até 20 dias úteis após assinatura em licitação; até 10 dias úteis em contratação direta.",
            "Lei nº 14.133/2021, art. 94.",
        ),
        (
            "Dados de obras contratadas",
            "Quantitativos e preços contratados em até 25 dias úteis da assinatura; executados e praticados em até 45 dias úteis da conclusão.",
            "Lei nº 14.133/2021, art. 94, §3º.",
        ),
    ]
    if not is_prefeitura:
        deadline_rows[4] = (
            "RREO consolidado do Município",
            "Até 30 dias após cada bimestre; a Câmara deve permitir rastrear seus dados que integram a consolidação.",
            "LRF, art. 52.",
        )
        deadline_rows[7] = (
            "Fiscalização da programação financeira",
            "O Executivo deve estabelecê-la em até 30 dias após a publicação do orçamento; a Câmara fiscaliza o cumprimento.",
            "LRF, arts. 8º e 59.",
        )
    add_legal_table(
        doc,
        ("INFORMAÇÃO OU ATO", "PRAZO / PERIODICIDADE", "BASE LEGAL"),
        deadline_rows,
        widths=(3.6, 8.2, 5.0),
    )

    doc.add_heading("8. Requisitos mínimos de entrega solicitados", level=1)
    add_numbered_items(
        doc,
        [
            (
                "Dados estruturados e completos. ",
                "Fornecer o conjunto integral em CSV, XLSX ou JSON quando a origem for uma base "
                "estruturada, além de eventual PDF de apresentação. Informar codificação, campos, "
                "tipos, unidades, chaves, período coberto e dicionário de dados.",
            ),
            (
                "Metadados e comprovação de atualidade. ",
                "Exibir fonte, unidade responsável, periodicidade, data e hora do registro, data e "
                "hora da publicação, última atualização e cobertura histórica.",
            ),
            (
                "Acesso integral ou parcialmente tarjado. ",
                "Se houver dado protegido, entregar a parte pública por certidão, extrato ou cópia "
                "com ocultação somente do trecho sigiloso, indicando a base legal específica da "
                "restrição, a autoridade classificadora, o prazo de sigilo e a via recursal.",
            ),
            (
                "Integridade, permanência e rastreabilidade. ",
                "Manter URLs permanentes, número de versão, anexos, histórico de substituições e "
                "identificador único que permita cruzar orçamento, licitação, contrato, empenho, "
                "liquidação, pagamento e prestação de contas.",
            ),
            (
                "Consulta sem barreiras. ",
                "Não exigir cadastro, senha, aplicativo proprietário ou aceite de rastreamento não "
                "essencial para acessar dados públicos; permitir pesquisa, filtros combinados, "
                "paginação transparente e exportação integral do resultado.",
            ),
            (
                "Calendário público de obrigações. ",
                "Publicar calendário com o prazo legal, data efetiva de publicação, link da entrega "
                "e responsável por RREO, RGF, audiências, contas, instrumentos orçamentários, "
                "contratos, obras e bases de execução em tempo real.",
            ),
            (
                "Adoção municipal da Lei de Governo Digital. ",
                "Informar se a Lei Federal nº 14.129/2021 foi adotada pelo Município por ato "
                "normativo próprio, nos termos de seu art. 2º, III, e encaminhar o respectivo ato. "
                "Na ausência, indicar a norma ou política local que rege dados abertos e serviços digitais.",
            ),
        ],
    )

    qualifier = doc.add_table(rows=1, cols=1)
    prevent_row_split(qualifier.rows[0])
    cell = qualifier.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_border(cell, left={"val": "single", "sz": "18", "color": BLUE})
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Critério de verificação sugerido. ")
    r.bold = True
    p.add_run(
        "Para cada obrigação, a resposta deve indicar: situação atual, URL direta, formato, "
        "período disponível, data da última atualização, unidade responsável, eventual fundamento "
        "de restrição e prazo de correção. Respostas genéricas como “disponível no portal”, sem "
        "link e sem identificação da base, não permitem verificar o atendimento."
    )


def add_sources(doc: Document, result_label: str, result_url: str) -> None:
    doc.add_heading("9. Referências utilizadas", level=1)
    sources = [
        (
            "Constituição da República Federativa do Brasil",
            CONSTITUTION_URL,
            "arts. 5º, 31, 37, 70 e 74",
        ),
        (
            "Lei nº 12.527/2011 (Lei de Acesso à Informação)",
            LAI_URL,
            "transparência ativa, forma de entrega, prazos e recursos",
        ),
        (
            "Lei Municipal nº 5.729/2013",
            LOCAL_LAI_URL,
            "regulamentação do acesso à informação em Varginha para Executivo e Legislativo",
        ),
        (
            "Lei Complementar nº 101/2000 (Lei de Responsabilidade Fiscal)",
            LRF_URL,
            "execução, relatórios, contas, audiências e fiscalização",
        ),
        (
            "Decreto nº 10.540/2020 (padrão do SIAFIC)",
            SIAFIC_URL,
            "tempo real, detalhamento e acesso público",
        ),
        (
            "Lei nº 14.129/2021 (Governo Digital)",
            GOV_DIGITAL_URL,
            "aplicação municipal condicionada ao art. 2º, III; padrões de dados abertos e reutilização",
        ),
        (
            "Lei nº 14.133/2021 (Licitações e Contratos)",
            PROCUREMENT_URL,
            "PNCP, editais, contratos, aditivos e obras",
        ),
        (
            "Lei nº 13.460/2017 (Direitos do Usuário e Ouvidoria)",
            USER_SERVICE_URL,
            "manifestações e prazo próprio de resposta",
        ),
        (
            "Lei nº 13.709/2018 (LGPD)",
            LGPD_URL,
            "tratamento de dados pessoais pelo poder público",
        ),
        (
            "Lei nº 13.146/2015 (Lei Brasileira de Inclusão)",
            ACCESSIBILITY_URL,
            "acessibilidade obrigatória nos sítios governamentais",
        ),
        (
            "Lei nº 8.159/1991 (Política Nacional de Arquivos)",
            ARCHIVES_URL,
            "gestão, preservação e acesso a documentos públicos",
        ),
        (
            "Cartilha PNTP 2025 — Atricon",
            PNTP_URL,
            "metodologia, matriz e níveis de transparência",
        ),
        (
            "Downloads oficiais do Radar da Transparência Pública",
            PNTP_DOWNLOADS_URL,
            "matrizes e materiais do ciclo vigente",
        ),
        (
            result_label,
            result_url,
            "painel público consultado; percentuais sujeitos à confirmação pelo órgão e pelo TCE-MG",
        ),
    ]
    for idx, (label, url, note) in enumerate(sources, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.2)
        p.paragraph_format.first_line_indent = Cm(-0.2)
        p.paragraph_format.space_after = Pt(3)
        p.add_run(f"{idx}. ").bold = True
        add_hyperlink(p, label, url)
        p.add_run(f" — {note}.")

    note = doc.add_table(rows=1, cols=1)
    cell = note.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_border(
        cell,
        left={"val": "single", "sz": "18", "color": BLUE},
    )
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Nota de cautela. ")
    r.bold = True
    p.add_run(
        "O painel público consultado apresenta a nota global e o percentual de critérios "
        "essenciais, mas não identifica, na página de resultado, cada item não atendido. "
        "Por isso, este requerimento solicita o espelho detalhado da avaliação e não afirma "
        "quais critérios específicos falharam."
    )


def add_signature(doc: Document) -> None:
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    closing.add_run(
        "O presente requerimento tem caráter cidadão, técnico e colaborativo, sem imputação "
        "prévia de irregularidade. O objetivo é ampliar a utilidade pública das informações, "
        "facilitar o controle social e apoiar a evolução do portal nos próximos ciclos do PNTP."
    )
    closing.paragraph_format.space_after = Pt(8)
    for line in [
        "_______________________________________________",
        "[NOME DO(A) REQUERENTE / REPRESENTANTE]",
        "[CPF ou CNPJ, se exigido pelo canal de protocolo]",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.size = Pt(9)
        if line.startswith("[NOME"):
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(NAVY)


def create_prefeitura() -> Path:
    doc = Document()
    configure_document(doc, "Prefeitura")
    add_cover(
        doc,
        [
            "À PREFEITURA MUNICIPAL DE VARGINHA",
            "À Controladoria-Geral do Município e ao Serviço de Informação ao Cidadão (e-SIC)",
            "Rua Júlio Paulo Marcellini, nº 50 • Varginha/MG • CEP 37018-050",
        ],
        "Melhorias no Portal da Transparência",
        "Acesso ao espelho da avaliação do PNTP e solicitação de plano de aprimoramento",
        "PNTP 2025 • Portal da Transparência do Poder Executivo Municipal",
    )

    doc.add_heading("1. Contexto e objetivo", level=1)
    add_body_text(
        doc,
        "O Programa Nacional de Transparência Pública (PNTP), coordenado pela Atricon com "
        "participação dos Tribunais de Contas, avalia a disponibilidade, a atualidade, a série "
        "histórica, a possibilidade de gravação de relatórios e os filtros de pesquisa dos "
        "portais públicos. A matriz de 2025 contém 176 critérios, dos quais 91 são aplicáveis "
        "ao Poder Executivo municipal: 11 essenciais, 59 obrigatórios e 21 recomendados."
    )
    add_body_text(
        doc,
        "No painel público consultado para a Prefeitura Municipal de Varginha constam índice "
        "geral de 81,8%, nível “Elevado”, e atendimento de 81,8% dos critérios essenciais. "
        "Como a página não apresenta o espelho critério a critério, esses dados são utilizados "
        "apenas como ponto de partida e devem ser confirmados pela Prefeitura e pelo TCE-MG."
    )
    add_summary_cards(
        doc,
        [
            ("81,8%", "ÍNDICE GERAL INFORMADO", BLUE),
            ("ELEVADO", "NÍVEL INFORMADO", NAVY),
            ("81,8%", "ESSENCIAIS INFORMADOS", RED),
        ],
    )

    doc.add_heading("2. Parte A — Pedido de acesso a informações existentes", level=1)
    add_body_text(
        doc,
        "Com fundamento nos arts. 10 e 11 da Lei nº 12.527/2011, solicita-se acesso às "
        "informações e aos documentos já existentes abaixo relacionados. Não se requer, nesta "
        "parte, a produção de estudo novo."
    )
    add_numbered_items(
        doc,
        [
            (
                "Espelho integral da avaliação. ",
                "Fornecer a autoavaliação e a validação do ciclo PNTP 2025, em formato "
                "eletrônico, com a resposta atribuída a cada critério aplicável, os links de "
                "evidência, as observações, as justificativas e, se existentes, os registros "
                "de revisão ou recurso.",
            ),
            (
                "Itens não atendidos ou parcialmente atendidos. ",
                "Informar, por critério, a classificação (essencial, obrigatório ou recomendado) "
                "e o componente de verificação não alcançado: disponibilidade, atualidade, série "
                "histórica, gravação de relatório ou filtro de pesquisa.",
            ),
            (
                "Responsáveis e providências já formalizadas. ",
                "Encaminhar atos, memorandos, ordens de serviço, cronogramas, contratos, chamados "
                "técnicos ou outros registros existentes relativos à correção dos itens apontados.",
            ),
            (
                "Documentos orçamentários e fiscais. ",
                "Informar os links permanentes e diretos para PPA, LDO e LOA, com as respectivas "
                "leis e anexos, bem como RGF, RREO, prestações de contas e versões simplificadas, "
                "organizados por exercício.",
            ),
            (
                "Dados de receita e despesa. ",
                "Informar os endereços de consulta e de exportação dos dados pormenorizados da "
                "receita e da despesa, incluindo previsão, lançamento e arrecadação; empenho, "
                "liquidação e pagamento; classificação orçamentária; favorecido; objeto; processo "
                "e procedimento licitatório, quando aplicável.",
            ),
        ],
    )

    doc.add_heading("3. Parte B — Solicitação colaborativa de providências", level=1)
    add_body_text(
        doc,
        "Sem confundir esta seção com o pedido de acesso da Parte A, solicita-se que a "
        "Administração avalie a adoção das providências abaixo e informe, na resposta, quais "
        "medidas serão acolhidas, a unidade responsável e o prazo estimado."
    )
    add_numbered_items(
        doc,
        [
            (
                "Criar uma área única de Planejamento e Orçamento. ",
                "Reunir PPA, LDO e LOA por exercício, distinguindo projeto, lei aprovada, "
                "alterações e anexos. Usar URLs estáveis, títulos padronizados, identificação "
                "clara do exercício e arquivos pesquisáveis e acessíveis.",
            ),
            (
                "Elevar a transparência da execução orçamentária. ",
                "Disponibilizar receita e despesa em tempo real, com detalhamento completo, "
                "metadados, data e hora da última atualização e ligação entre empenho, liquidação, "
                "pagamento, contrato, licitação e fornecedor.",
            ),
            (
                "Oferecer dados reutilizáveis. ",
                "Permitir filtros combináveis, busca textual, exportação integral do resultado "
                "em CSV/XLSX e, quando viável, acesso por API, preservando série histórica mínima "
                "de três exercícios e evitando relatórios limitados apenas ao que aparece na tela.",
            ),
            (
                "Publicar painel de conformidade do PNTP. ",
                "Disponibilizar o espelho da avaliação, as evidências, o status de cada correção "
                "e um plano de ação com responsáveis e prazos, mantendo registro das melhorias "
                "entre ciclos.",
            ),
            (
                "Melhorar acessibilidade e experiência de uso. ",
                "Revisar contraste, redimensionamento de texto, navegação por teclado, rótulos "
                "para leitores de tela, caminho de navegação, mapa do site, linguagem simples, "
                "SIC, Ouvidoria, política de privacidade e canais de contato.",
            ),
        ],
    )

    doc.add_heading("4. Priorização sugerida", level=1)
    add_priority_table(
        doc,
        [
            (
                "IMEDIATA",
                "Corrigir todos os critérios essenciais e links quebrados",
                "Alcançar 100% dos essenciais, condição necessária para os selos Prata, Ouro e Diamante",
            ),
            (
                "30–60 DIAS",
                "Centralizar PPA, LDO, LOA, RGF, RREO, receitas e despesas",
                "Reduzir dispersão, facilitar localização e comprovação das evidências",
            ),
            (
                "60–90 DIAS",
                "Filtros, exportações, série histórica, acessibilidade e metadados",
                "Elevar a nota e tornar o portal efetivamente utilizável por cidadãos e sistemas",
            ),
            (
                "CONTÍNUA",
                "Painel PNTP, rotina de atualização e testes periódicos",
                "Evitar regressões e manter prontidão para o ciclo vigente",
            ),
        ],
    )

    doc.add_heading("5. Forma e prazo da resposta", level=1)
    add_body_text(
        doc,
        "Solicita-se resposta eletrônica, preferencialmente em planilha aberta para o espelho "
        "critério a critério, acompanhada dos documentos e links diretos. Nos termos do art. 11 "
        "da Lei nº 12.527/2011, a informação disponível deve ser fornecida imediatamente; não "
        "sendo possível, o prazo é de até 20 dias, prorrogável por mais 10 dias mediante "
        "justificativa expressa e ciência do requerente."
    )
    add_body_text(
        doc,
        "Caso alguma informação seja de competência de outro órgão, solicita-se a indicação "
        "de quem a detém ou a remessa do pedido, com ciência ao requerente. Eventual negativa "
        "total ou parcial deve ser fundamentada e acompanhada da orientação sobre recurso."
    )
    add_legal_framework(doc, is_prefeitura=True)
    add_sources(
        doc,
        "Resultado público consultado — Prefeitura Municipal de Varginha",
        PREF_RESULT_URL,
    )
    add_signature(doc)

    path = OUT_DIR / "Requerimento_Melhorias_Portal_Transparencia_Prefeitura_Varginha.docx"
    doc.save(path)
    return path


def create_camara() -> Path:
    doc = Document()
    configure_document(doc, "Câmara Municipal")
    add_cover(
        doc,
        [
            "À CÂMARA MUNICIPAL DE VARGINHA",
            "À Controladoria Interna, à Ouvidoria e ao Serviço de Informação ao Cidadão (e-SIC)",
            "Praça Governador Benedito Valadares, nº 11 • Centro • Varginha/MG • CEP 37002-020",
        ],
        "Melhorias no Portal da Transparência",
        "Acesso ao espelho da avaliação do PNTP e solicitação de plano de aprimoramento",
        "PNTP 2025 • Portal da Transparência do Poder Legislativo Municipal",
    )

    doc.add_heading("1. Contexto e objetivo", level=1)
    add_body_text(
        doc,
        "O Programa Nacional de Transparência Pública (PNTP), coordenado pela Atricon com "
        "participação dos Tribunais de Contas, avalia a disponibilidade, a atualidade, a série "
        "histórica, a possibilidade de gravação de relatórios e os filtros de pesquisa dos "
        "portais públicos. A matriz de 2025 contém 176 critérios, dos quais 82 são aplicáveis "
        "ao Poder Legislativo municipal: 6 essenciais, 60 obrigatórios e 16 recomendados."
    )
    add_body_text(
        doc,
        "No painel público consultado para a Câmara Municipal de Varginha constam índice geral "
        "de 63,3%, nível “Intermediário”, atendimento de 85,7% dos critérios essenciais e "
        "variação positiva de 21,4 pontos percentuais. A evolução é relevante, mas a página não "
        "apresenta o espelho critério a critério; por isso, os dados devem ser confirmados pela "
        "Câmara e pelo TCE-MG."
    )
    add_summary_cards(
        doc,
        [
            ("63,3%", "ÍNDICE GERAL INFORMADO", BLUE),
            ("INTERMEDIÁRIO", "NÍVEL INFORMADO", NAVY),
            ("85,7%", "ESSENCIAIS INFORMADOS", RED),
            ("+21,4 p.p.", "EVOLUÇÃO INFORMADA", GREEN),
        ],
    )

    doc.add_heading("2. Parte A — Pedido de acesso a informações existentes", level=1)
    add_body_text(
        doc,
        "Com fundamento nos arts. 10 e 11 da Lei nº 12.527/2011, solicita-se acesso às "
        "informações e aos documentos já existentes abaixo relacionados. Não se requer, nesta "
        "parte, a produção de estudo novo."
    )
    add_numbered_items(
        doc,
        [
            (
                "Espelho integral da avaliação. ",
                "Fornecer a autoavaliação e a validação do ciclo PNTP 2025, em formato "
                "eletrônico, com a resposta de cada critério aplicável, links de evidência, "
                "observações, justificativas e, se existentes, registros de revisão ou recurso.",
            ),
            (
                "Itens não atendidos ou parcialmente atendidos. ",
                "Informar, por critério, a classificação (essencial, obrigatório ou recomendado) "
                "e o componente não alcançado: disponibilidade, atualidade, série histórica, "
                "gravação de relatório ou filtro de pesquisa.",
            ),
            (
                "Responsáveis e providências já formalizadas. ",
                "Encaminhar atos, memorandos, ordens de serviço, cronogramas, contratos, chamados "
                "técnicos ou outros registros existentes relativos à correção dos apontamentos.",
            ),
            (
                "Execução orçamentária e gestão fiscal. ",
                "Informar os links diretos para receitas; despesas empenhadas, liquidadas e pagas; "
                "classificação orçamentária; favorecidos; objetos; processos; procedimentos "
                "licitatórios; e Relatórios de Gestão Fiscal, com histórico e exportação.",
            ),
            (
                "Atividade legislativa e administrativa. ",
                "Informar os links e bases de dados que reúnem pautas, atas, presenças, votações "
                "nominais, projetos e tramitações, leis e atos, trabalhos das comissões, contas "
                "do Chefe do Executivo, contratos, licitações, pessoal, remuneração e diárias.",
            ),
        ],
    )

    doc.add_heading("3. Parte B — Solicitação colaborativa de providências", level=1)
    add_body_text(
        doc,
        "Sem confundir esta seção com o pedido de acesso da Parte A, solicita-se que a Câmara "
        "avalie as providências abaixo e informe quais medidas serão acolhidas, a unidade "
        "responsável e o prazo estimado."
    )
    add_numbered_items(
        doc,
        [
            (
                "Corrigir prioritariamente os critérios essenciais. ",
                "Garantir portal próprio ou compartilhado claramente acessível, receitas, despesas "
                "empenhadas/liquidadas/pagas com classificação adequada e RGF, com evidências "
                "estáveis e atualizadas.",
            ),
            (
                "Consolidar a transparência legislativa. ",
                "Organizar composição e biografias; leis e atos; projetos com ementa, anexos, "
                "situação, autoria, relatoria e tramitação; pautas do Plenário e das Comissões; "
                "atas e presença; votações nominais; julgamento das contas do Executivo; "
                "transmissões; verbas indenizatórias e atividades parlamentares.",
            ),
            (
                "Detalhar pessoal e remuneração. ",
                "Oferecer relação nominal, cargo ou função, lotação, admissão, exoneração ou "
                "inativação, carga horária, remuneração individualizada e tabela remuneratória, "
                "com busca, filtros e histórico, observadas as salvaguardas legais aplicáveis.",
            ),
            (
                "Integrar licitações e contratos. ",
                "Publicar relação sequencial, editais e documentos das fases interna e externa, "
                "dispensas e inexigibilidades, atas de adesão, contratos, aditivos, fiscais, "
                "fornecedores e sanções, com ligação entre todos os registros.",
            ),
            (
                "Padronizar diárias e deslocamentos. ",
                "Divulgar beneficiário, cargo, valor total, quantidade, período, motivo, destino "
                "e tabela de valores, com filtros, exportação e série histórica.",
            ),
            (
                "Aprimorar usabilidade, dados abertos e acessibilidade. ",
                "Incluir busca textual, filtros combináveis, exportação integral em CSV/XLSX, "
                "metadados, data da atualização, URLs permanentes, navegação por teclado, alto "
                "contraste, redimensionamento, caminho de páginas, mapa do site, SIC, Ouvidoria "
                "e política de privacidade.",
            ),
            (
                "Publicar plano de evolução do PNTP. ",
                "Disponibilizar o espelho da avaliação, evidências, status de correção, responsáveis "
                "e prazos, com acompanhamento público entre os ciclos.",
            ),
        ],
    )

    doc.add_heading("4. Priorização sugerida", level=1)
    add_priority_table(
        doc,
        [
            (
                "IMEDIATA",
                "Corrigir todos os critérios essenciais e links indisponíveis",
                "Alcançar 100% dos essenciais, condição necessária para os selos Prata, Ouro e Diamante",
            ),
            (
                "30–60 DIAS",
                "Receita, despesa, RGF, pessoal, licitações e contratos",
                "Concentrar informações de maior peso e obrigação legal",
            ),
            (
                "60–90 DIAS",
                "Atividade legislativa, diárias, filtros, exportações e acessibilidade",
                "Elevar a nota global e facilitar o acompanhamento do mandato e dos gastos",
            ),
            (
                "CONTÍNUA",
                "Painel PNTP, rotina de atualização e testes periódicos",
                "Evitar regressões e manter prontidão para o ciclo vigente",
            ),
        ],
    )

    doc.add_heading("5. Forma e prazo da resposta", level=1)
    add_body_text(
        doc,
        "Solicita-se resposta eletrônica, preferencialmente em planilha aberta para o espelho "
        "critério a critério, acompanhada dos documentos e links diretos. Nos termos do art. 11 "
        "da Lei nº 12.527/2011, a informação disponível deve ser fornecida imediatamente; não "
        "sendo possível, o prazo é de até 20 dias, prorrogável por mais 10 dias mediante "
        "justificativa expressa e ciência do requerente."
    )
    add_body_text(
        doc,
        "Caso alguma informação seja de competência de outro órgão, solicita-se a indicação "
        "de quem a detém ou a remessa do pedido, com ciência ao requerente. Eventual negativa "
        "total ou parcial deve ser fundamentada e acompanhada da orientação sobre recurso."
    )
    add_legal_framework(doc, is_prefeitura=False)
    add_sources(
        doc,
        "Resultado público consultado — Câmara Municipal de Varginha",
        CAM_RESULT_URL,
    )
    add_signature(doc)

    path = OUT_DIR / "Requerimento_Melhorias_Portal_Transparencia_Camara_Varginha.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    for output in (create_prefeitura(), create_camara()):
        print(output)
