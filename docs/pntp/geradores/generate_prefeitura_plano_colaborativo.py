from __future__ import annotations

import csv
import re
from collections import Counter, OrderedDict
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

import generate_camara_plano_colaborativo as base


ROOT = Path(__file__).resolve().parents[1]
CSV_PATH = ROOT / "docs" / "pntp" / "criterios-pntp-2026.csv"
OUT_DIR = ROOT / "docs" / "pntp" / "solicitacoes-portal"
OUT_PATH = OUT_DIR / "Plano_Colaborativo_Melhorias_Portal_Transparencia_Prefeitura_Varginha.docx"

PORTAL_URL = "https://www.varginha.mg.gov.br/portal/transparencia"
OPEN_DATA_URL = "https://www.varginha.mg.gov.br/portal/dados-abertos"
GAZETTE_URL = "https://www.varginha.mg.gov.br/portal/diario-oficial/"
RESULT_URL = "https://radartransparente.com.br/avaliacao/59f001b0-a3d0-42cb-953b-0d02db067387"


def load_criteria() -> list[dict[str, str]]:
    allowed = {
        "COMUM",
        "COMUM (EXCETO ESTATAIS INDEPENDENTES)",
        "COMUM (EXCETO ESTATAIS)",
        "EXECUTIVO",
        "EXECUTIVO e CONSÓRCIOS",
    }
    with CSV_PATH.open(encoding="utf-8-sig", newline="") as stream:
        rows = [row for row in csv.DictReader(stream) if row["matriz"] in allowed]
    rows.sort(key=lambda row: tuple(int(number) for number in re.findall(r"\d+", row["id"])))
    assert len(rows) == 96, f"Esperados 96 critérios; encontrados {len(rows)}"
    return rows


def configure_document(doc: Document) -> None:
    base.configure_document(doc)
    header = doc.sections[0].header.tables[0]
    header.cell(0, 1).text = "PREFEITURA"
    p = header.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        base.set_run_font(run, 9, base.BLUE, bold=True)


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("PORTAL DA TRANSPARÊNCIA")
    base.set_run_font(r, 10, base.BLUE, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Prefeitura Municipal de Varginha")
    base.set_run_font(r, 25, base.NAVY, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Diagnóstico de falhas e plano colaborativo de regularização integral em 60 dias")
    base.set_run_font(r, 13, base.BLUE, bold=True)

    rows = [
        ("DESTINATÁRIO", "Prefeitura Municipal de Varginha — Poder Executivo"),
        ("PROPONENTE", "Fiscaliza Varginha — iniciativa cidadã independente"),
        ("DATA-BASE", "25 de julho de 2026"),
        ("FINALIDADE", "Apoiar a correção das pendências, o cumprimento dos prazos legais e a melhoria da transparência ativa."),
    ]
    base.add_table(doc, ["IDENTIFICAÇÃO", "INFORMAÇÃO"], rows, [2100, 7538], 9.3, True)
    base.add_callout(
        doc,
        "Tom e finalidade",
        "O documento reconhece os avanços já existentes, registra falhas, deficiências e atrasos que "
        "precisam ser enfrentados e propõe cooperação técnica para concluir a regularização em 60 dias. "
        "Não presume dolo ou irregularidade e não substitui a avaliação oficial do Tribunal de Contas.",
        fill=base.PALE_BLUE,
        accent=base.BLUE,
    )


def add_status_strip(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    values = [
        ("81,8%", "ÍNDICE INFORMADO", base.BLUE),
        ("ELEVADO", "NÍVEL INFORMADO", base.NAVY),
        ("81,8%", "ESSENCIAIS INFORMADOS", base.RED),
        ("60 DIAS", "META DE CONCLUSÃO", base.GREEN),
    ]
    base.set_table_geometry(table, [2409, 2409, 2409, 2411], indent_dxa=0)
    for cell, (value, label, fill) in zip(table.rows[0].cells, values):
        base.set_cell_shading(cell, fill)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        base.set_run_font(r, 14 if len(value) > 8 else 16, base.WHITE, bold=True)
        r = p.add_run(f"\n{label}")
        base.set_run_font(r, 7.5, base.WHITE, bold=True)
    base.prevent_row_split(table.rows[0])


def action_for(row: dict[str, str]) -> str:
    special = {
        "3.2": "Detalhar a receita por categoria econômica, origem, espécie, desdobramento e fonte, com filtros e exportação.",
        "3.3": "Publicar a dívida ativa em consulta pesquisável e exportável, com atualização, natureza, situação e proteção dos dados legalmente restritos.",
        "11.4": "Reunir parecer prévio do TCE, decreto legislativo, ata e resultado do julgamento das contas do Prefeito.",
        "11.6": "Publicar o RREO e todos os anexos até 30 dias após cada bimestre, preservando assinatura e série histórica.",
        "11.8": "Centralizar o PPA vigente e anteriores, lei, anexos, programas, metas, revisões e versões consolidadas.",
        "11.9": "Centralizar a LDO de cada exercício, lei, anexos fiscais, riscos, metas, alterações e versão consolidada.",
        "11.10": "Centralizar a LOA de cada exercício, lei, quadros, anexos, créditos adicionais e execução vinculada.",
        "16.1": "Publicar cada renúncia ou benefício tributário com fundamento legal, tributo, modalidade, vigência e estimativa.",
        "16.2": "Divulgar previsão e realização das renúncias fiscais, metodologia, memória de cálculo e medidas de compensação.",
        "16.3": "Disponibilizar beneficiários, valores e benefícios concedidos, observadas apenas as restrições legais específicas.",
        "16.4": "Publicar projetos de incentivo à cultura e ao esporte, proponente, objeto, valor autorizado, captado e executado.",
        "17.1": "Publicar emendas parlamentares recebidas por autor, origem, objeto, valor, instrumento e unidade executora.",
        "17.2": "Demonstrar execução orçamentária e financeira de cada emenda, inclusive empenho, pagamento e saldo.",
        "17.3": "Divulgar transferências especiais, plano de trabalho, conta, execução, beneficiário final e prestação de contas.",
        "18.1": "Publicar Plano Municipal de Saúde, programação anual, relatórios de gestão, metas, resultados e deliberações do Conselho.",
        "18.2": "Manter relação atualizada de unidades, serviços, horários, profissionais, especialidades e formas de acesso.",
        "18.3": "Publicar listas de espera por serviço, posição ou tempo estimado, critérios de priorização e data de atualização, preservando a identidade do paciente.",
        "18.4": "Divulgar estoque de medicamentos por unidade, item, disponibilidade e data de atualização.",
        "18.5": "Publicar repasses, despesas, contratos e indicadores da saúde em painel pesquisável e exportável.",
        "18.6": "Dar transparência à composição, agenda, atas, resoluções e deliberações do Conselho Municipal de Saúde.",
        "19.1": "Publicar Plano Municipal de Educação, metas, indicadores, monitoramento e relatórios de execução.",
        "19.2": "Divulgar vagas e listas de espera da educação infantil com critérios, atualização e proteção dos dados das crianças.",
        "19.3": "Publicar composição, agenda, atas, pareceres e deliberações dos conselhos de educação e controle social.",
        "19.4": "Divulgar serviços, critérios, vagas, filas, benefícios e resultados da assistência social, com atualização e proteção de dados.",
    }
    key = row["id"].rstrip(".")
    return special.get(key, base.action_for(row))


def evidence_for(row: dict[str, str]) -> str:
    dimension = row["dimensao"]
    if dimension == "Renúncias de Receita":
        return "Base legal; tributo; beneficiário quando cabível; modalidade; valor previsto/realizado; vigência; CSV."
    if dimension == "Emendas Parlamentares":
        return "Autor/origem; objeto; valor; plano; instrumento; empenho; pagamento; saldo; prestação; CSV."
    if dimension == "Saúde":
        return "Documento/painel integral; unidade; data da atualização; série; indicador; filtro; exportação; proteção do paciente."
    if dimension == "Educação e Assistência Social":
        return "Plano/serviço; meta ou fila; critérios; resultado; conselho; data de atualização; exportação."
    return base.evidence_for(row)


def add_main_content(doc: Document, criteria: list[dict[str, str]]) -> None:
    counts = Counter(row["classe"] for row in criteria)
    doc.add_heading("1. Síntese executiva", level=1)
    base.add_body(
        doc,
        "O resultado público anteriormente consultado para a Prefeitura indicou índice geral de 81,8%, "
        "classificação “Elevado” e atendimento de 81,8% dos critérios essenciais. O desempenho demonstra "
        "uma base relevante, mas ainda não representa atendimento integral: critérios essenciais abaixo "
        "de 100%, publicações atrasadas e lacunas de qualidade impedem a conclusão plena do ciclo."
    )
    add_status_strip(doc)
    base.add_body(
        doc,
        "A Prefeitura mantém portal institucional, área de dados abertos, Diário Oficial e sistemas "
        "temáticos. O desafio é tornar esse conjunto uma experiência única e verificável, com conteúdo "
        "completo, atualizado, pesquisável, exportável, acessível e ligado aos respectivos atos, processos "
        "e registros contábeis."
    )
    base.add_callout(
        doc,
        "Escopo completo",
        f"A matriz PNTP 2026 reúne 96 critérios aplicáveis ao Executivo municipal: {counts['Essencial']} "
        f"essenciais, {counts['Obrigatória']} obrigatórios e {counts['Recomendada']} recomendados. "
        "O Anexo A converte todos eles em ações e evidências de entrega.",
        fill=base.PALE_GOLD,
        accent=base.GOLD,
    )
    base.add_callout(
        doc,
        "Meta de 60 dias",
        "Até o encerramento do período, os 96 critérios deverão estar validados; o passivo de publicações "
        "deverá estar regularizado; e cada item aplicável deverá possuir URL, evidência, responsável, data "
        "de atualização e teste final registrados.",
        fill=base.PALE_GREEN,
        accent=base.GREEN,
    )

    doc.add_heading("2. Como os apontamentos foram classificados", level=1)
    base.add_table(
        doc, ["CLASSIFICAÇÃO", "SIGNIFICADO", "TRATAMENTO"],
        [
            ["Falha confirmada", "Resultado público ou ausência objetivamente verificável.", "Corrigir, publicar evidência e registrar a data."],
            ["Deficiência observada", "Barreira de navegação, organização ou forma de entrega.", "Validar com o setor e ajustar a experiência."],
            ["Atraso relatado", "Entrega informada como pendente ou fora da rotina esperada.", "Inventariar, atribuir responsável e recuperar o passivo."],
            ["Item a validar", "Critério sem espelho detalhado público disponível.", "Conferir sem presumir falha antes da validação."],
        ], [2100, 3500, 4038], 9.0, True
    )

    doc.add_heading("3. Apontamento das principais falhas e deficiências", level=1)
    findings = [
        ["F-01", "Falha confirmada", "Índice geral informado de 81,8%, ainda sem atendimento integral.", "Publicar plano critério a critério e evidenciar a correção dos itens restantes.", "Imediata"],
        ["F-02", "Falha confirmada", "Essenciais informados em 81,8%, abaixo dos 100% necessários ao selo.", "Priorizar todos os essenciais não atendidos ou parciais.", "Imediata"],
        ["F-03", "Atraso relatado", "Há entregas e publicações em atraso, sem inventário público consolidado.", "Criar força-tarefa, listar o passivo por competência e publicar calendário.", "0-30 dias"],
        ["F-04", "Deficiência observada", "Conteúdo distribuído entre portal, dados abertos, Diário Oficial, sistemas, legislação e PNCP.", "Criar página única por tema com links diretos, contexto, período e responsável.", "0-30 dias"],
        ["F-05", "Deficiência observada", "Não há painel único de atualidade, cobertura histórica e responsável por base.", "Exibir última atualização, competência, formato, setor e próxima carga.", "0-30 dias"],
        ["F-06", "Item a validar", "Não está disponível um espelho público PNTP critério a critério.", "Publicar matriz com status, URL, evidência, responsável e revisão.", "0-30 dias"],
        ["F-07", "Item a validar", "PPA, LDO, LOA, RREO, RGF e anexos exigem centralização, versões e vínculos.", "Criar repositório orçamentário por ciclo, exercício e período.", "0-30 dias"],
        ["F-08", "Item a validar", "Saúde, educação, assistência, filas e medicamentos exigem atualização frequente e data clara.", "Definir carga, indicador de atualidade, histórico e proteção adequada.", "0-45 dias"],
        ["F-09", "Item a validar", "Renúncias e emendas precisam permitir rastrear fundamento, beneficiário, objeto e execução.", "Publicar bases abertas e relacionar previsão, concessão e execução.", "0-45 dias"],
        ["F-10", "Deficiência de governança", "Não está visível rotina pública única de responsáveis, prazos e testes.", "Instituir comitê, calendário, alertas, painel e revisão mensal.", "0-10 dias"],
    ]
    base.add_table(doc, ["ID", "EVIDÊNCIA", "APONTAMENTO", "MELHORIA COLABORATIVA", "PRAZO"],
                   findings, [620, 1500, 2580, 3938, 1000], 8.2, True)
    base.add_callout(
        doc, "Cautela metodológica",
        "Como o espelho detalhado da avaliação não estava disponível na página pública consultada, "
        "não se atribui automaticamente cada critério como falha. A Prefeitura poderá validar o anexo, "
        "identificar o que já atende e concentrar a correção nos itens parciais, atrasados ou ausentes.",
        fill=base.PALE_RED, accent=base.RED
    )

    doc.add_heading("4. Plano de regularização integral em 60 dias", level=1)
    roadmap = [
        ["Dias 1-10", "Governança e diagnóstico", "Nomear ponto focal; reunir setores; validar 96 critérios; listar atrasos, links, documentos e dependências.", "100% classificados, com responsável e prazo."],
        ["Dias 11-20", "Essenciais e fiscal", "Corrigir essenciais; completar receita, despesa, dívida ativa, RREO, RGF, PPA, LDO, LOA e contas.", "100% dos essenciais com evidência testada."],
        ["Dias 21-30", "Passivo de publicações", "Regularizar documentos vencidos de pessoal, diárias, compras, contratos, obras, convênios e prestação de contas.", "Passivo integralmente publicado e datado."],
        ["Dias 31-40", "Dados e atendimento", "Implantar busca, filtros, exportação, metadados, SIC, Ouvidoria, Carta de Serviços, LGPD e acessibilidade.", "Bases abertas, pesquisáveis e acessíveis."],
        ["Dias 41-50", "Políticas finalísticas", "Completar saúde, educação, assistência, renúncias, emendas, conselhos, filas e medicamentos.", "Políticas e recursos rastreáveis."],
        ["Dias 51-60", "Teste e entrega", "Testar 96 critérios, corrigir remanescentes e publicar matriz final, evidências, histórico e relatório.", "Nenhuma pendência aplicável aberta."],
        ["Após o dia 60", "Manutenção", "Revisão mensal, alertas de prazo, testes trimestrais e registro público de regressões.", "Impedir novos atrasos."],
    ]
    base.add_table(doc, ["JANELA", "FOCO", "ENTREGAS", "RESULTADO ESPERADO"],
                   roadmap, [1250, 1800, 4438, 2150], 8.7, True)
    base.add_callout(
        doc, "Regra de conclusão",
        "Uma ação somente será concluída com publicação acessível, URL direta, conteúdo integral, data "
        "de atualização, responsável e teste registrado. Dependência de fornecedor deverá ter solução "
        "alternativa e não poderá manter informação indisponível após o 60º dia.",
        fill=base.PALE_GREEN, accent=base.GREEN
    )

    doc.add_heading("5. Pacotes de melhoria por área", level=1)
    packages = [
        ["Governança e arquitetura", "Matriz pública, responsáveis, calendário, páginas temáticas, URLs permanentes, busca e alertas."],
        ["Receita, despesa e dívida", "Previsão/realização; classificação; empenho/liquidação/pagamento; credor; processo; dívida ativa; exportação."],
        ["Planejamento e contas", "PPA, LDO, LOA, RREO, RGF, balanços, metas, audiências, decisões e séries históricas."],
        ["Compras, contratos e obras", "Processo integral, PNCP, aditivos, fiscais, execução, medições, pagamentos, fotos e paralisações."],
        ["Pessoal e diárias", "Base nominal, remuneração, lotação, vínculos, atos, concursos, terceirizados, viagens e exportação."],
        ["Convênios e transferências", "Instrumento, partes, objeto, valores, vigência, execução, prestação e inteiro teor."],
        ["Renúncias e emendas", "Base legal, beneficiários, valores, autores, objetos, planos, empenhos, pagamentos e resultados."],
        ["Saúde", "Plano, serviços, profissionais, listas de espera, critérios, medicamentos, despesas, indicadores e conselho."],
        ["Educação e assistência", "Planos, metas, vagas, listas, critérios, serviços, benefícios, resultados e conselhos."],
        ["SIC, Ouvidoria e Carta", "Canais simples, protocolos, prazos, recursos, estatísticas, serviços e acompanhamento."],
        ["Dados abertos e integração", "CSV/XLSX/JSON, dicionário, metadados, APIs e chaves entre processo, contrato, empenho e pagamento."],
        ["Acessibilidade, LGPD e participação", "Teclado, leitor de tela, contraste, linguagem simples, acesso parcial, audiências e devolutivas."],
    ]
    base.add_table(doc, ["ÁREA", "MELHORIAS PROPOSTAS"], packages, [2550, 7088], 9.0, True)

    doc.add_heading("6. Padrão mínimo para cada publicação", level=1)
    standards = [
        ["Identificação", "Título, unidade, responsável, exercício/competência e descrição."],
        ["Atualidade", "Data do registro, publicação, última atualização e periodicidade."],
        ["Completude", "Conteúdo integral e anexos; declaração negativa quando o evento não ocorreu."],
        ["Formato", "Consulta e exportação integral em CSV, XLSX ou JSON; PDF como complemento."],
        ["Pesquisa", "Busca textual, filtros combináveis, ordenação e paginação transparente."],
        ["Histórico", "Ao menos três exercícios ou período legal; versões anteriores preservadas."],
        ["Integração", "Chaves entre processo, licitação, contrato, empenho, pagamento e PNCP."],
        ["Acessibilidade", "Teclado, leitor de tela, contraste, redimensionamento e documentos pesquisáveis."],
        ["Rastreabilidade", "URL permanente, fonte, versão, dicionário de dados e contato técnico."],
        ["Proteção de dados", "Ocultar apenas o trecho protegido, fornecer a parte pública e indicar a base legal."],
    ]
    base.add_table(doc, ["ELEMENTO", "PADRÃO DE ENTREGA"], standards, [2100, 7538], 9.0, True)

    doc.add_heading("7. Prazos e periodicidades prioritários", level=1)
    deadlines = [
        ["Execução orçamentária e financeira", "Até o primeiro dia útil após o registro contábil.", "Decreto nº 10.540/2020."],
        ["Pedido de acesso", "Imediato; se inviável, 20 dias + uma prorrogação de 10, justificada.", "LAI e Lei Municipal nº 5.729/2013."],
        ["Ouvidoria", "30 dias + uma prorrogação de 30, justificada.", "Lei nº 13.460/2017."],
        ["RREO", "Até 30 dias após cada bimestre.", "LRF, arts. 52 e 53."],
        ["RGF", "Até 30 dias após cada quadrimestre.", "LRF, arts. 54 e 55."],
        ["Metas fiscais", "Audiências até o fim de fevereiro, maio e setembro.", "LRF, art. 9º, §4º."],
        ["Programação financeira", "Até 30 dias após a publicação do orçamento.", "LRF, art. 8º."],
        ["Contratos no PNCP", "20 dias úteis após licitação; 10 dias úteis na contratação direta.", "Lei nº 14.133/2021, art. 94."],
        ["Obras", "Quantidades/preços em 25 dias úteis; executados/pagos em 45 dias úteis da conclusão.", "Lei nº 14.133/2021, art. 94, §3º."],
    ]
    base.add_table(doc, ["INFORMAÇÃO", "PRAZO", "BASE"], deadlines, [2600, 4388, 2650], 8.7, True)

    doc.add_heading("8. Governança participativa", level=1)
    governance = [
        ["Gabinete e Secretaria de Governo", "Patrocinar o plano, remover impedimentos e acompanhar marcos."],
        ["Controladoria", "Validar critérios, evidências, prazos e correções."],
        ["Fazenda e Contabilidade", "Atualizar receita, despesa, dívida, relatórios, orçamento e contas."],
        ["Planejamento", "Coordenar PPA, LDO, LOA, metas, indicadores e audiências."],
        ["Administração e RH", "Atualizar pessoal, remuneração, cargos, concursos, terceirizados e diárias."],
        ["Compras, Contratos e Obras", "Publicar processos, PNCP, instrumentos, execução, medições e pagamentos."],
        ["Saúde, Educação e Assistência", "Atualizar serviços, filas, estoques, metas, resultados e conselhos."],
        ["SIC, Ouvidoria e LGPD", "Garantir canais, prazos, relatórios, Carta, privacidade e acesso parcial."],
        ["Tecnologia e Comunicação", "Garantir disponibilidade, integração, acessibilidade, dados abertos e linguagem cidadã."],
        ["Sociedade civil", "Testar o portal, indicar barreiras e acompanhar a matriz pública."],
    ]
    base.add_table(doc, ["RESPONSÁVEL", "CONTRIBUIÇÃO SUGERIDA"], governance, [2700, 6938], 9.0, True)
    base.add_callout(
        doc, "Proposta de participação",
        "Sugere-se reunião técnica inicial e devolutiva pública no 30º e no 60º dia. O Fiscaliza Varginha "
        "pode contribuir com testes de navegação, organização das evidências e comunicação em linguagem "
        "cidadã, respeitada a autonomia administrativa da Prefeitura.",
        fill=base.PALE_GREEN, accent=base.GREEN
    )

    doc.add_heading("9. Painel público de acompanhamento", level=1)
    base.add_body(doc, "O painel deve conter uma linha por critério e permanecer público entre os ciclos de avaliação.")
    tracker = [
        ["Critério PNTP", "ID, classe, matriz e descrição."],
        ["Situação", "Atende; parcial; não atende; não se aplica; em validação."],
        ["Evidência", "URL direta, arquivo, tela ou relatório comprobatório."],
        ["Qualidade", "Disponibilidade, atualidade, histórico, exportação e pesquisa."],
        ["Responsável", "Secretaria, setor e agente encarregado."],
        ["Prazo", "Data prevista e data efetiva."],
        ["Observação", "Pendência, justificativa, dependência e próxima ação."],
        ["Revisão", "Data, responsável e resultado do último teste."],
    ]
    base.add_table(doc, ["CAMPO", "CONTEÚDO"], tracker, [2200, 7438], 9.0, True)

    doc.add_heading("10. Base normativa e referências", level=1)
    references = [
        ("Portal da Transparência da Prefeitura", PORTAL_URL, "acesso institucional"),
        ("Dados Abertos da Prefeitura", OPEN_DATA_URL, "bases e arquivos públicos"),
        ("Diário Oficial do Município", GAZETTE_URL, "publicidade oficial"),
        ("Downloads oficiais PNTP", base.PNTP_DOWNLOADS_URL, "matriz do ciclo 2026"),
        ("Cartilha PNTP — Atricon", base.PNTP_CARTILHA_URL, "metodologia"),
        ("Resultado público anteriormente consultado", RESULT_URL, "índice sujeito à confirmação oficial"),
        ("Lei nº 12.527/2011 — LAI", base.LAI_URL, "transparência ativa, formatos e prazos"),
        ("Lei Municipal nº 5.729/2013", base.LOCAL_LAI_URL, "regulamentação local"),
        ("Lei Complementar nº 101/2000 — LRF", base.LRF_URL, "relatórios, contas e fiscalização"),
        ("Decreto nº 10.540/2020 — SIAFIC", base.SIAFIC_URL, "tempo real e detalhamento"),
        ("Lei nº 14.133/2021", base.PROCUREMENT_URL, "PNCP, contratos e obras"),
        ("Lei nº 13.460/2017", base.OUVIDORIA_URL, "Ouvidoria e direitos do usuário"),
        ("Lei nº 13.709/2018 — LGPD", base.LGPD_URL, "proteção de dados no poder público"),
        ("Lei nº 13.146/2015 — LBI", base.LBI_URL, "acessibilidade digital"),
    ]
    for idx, (label, url, note) in enumerate(references, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.15)
        p.paragraph_format.first_line_indent = Cm(-0.15)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{idx}. ")
        base.set_run_font(r, 9.8, base.BLUE, bold=True)
        base.add_hyperlink(p, label, url)
        r = p.add_run(f" — {note}.")
        base.set_run_font(r, 9.8, base.INK)


def add_annex(doc: Document, criteria: list[dict[str, str]]) -> None:
    doc.add_heading("Anexo A — Checklist completo PNTP 2026 para a Prefeitura", level=1)
    base.add_body(
        doc,
        "O anexo reúne os 96 critérios das matrizes COMUM, COMUM (exceto estatais), COMUM (exceto "
        "estatais independentes), EXECUTIVO e EXECUTIVO e CONSÓRCIOS. A situação não é presumida: "
        "cada unidade deve validar, regularizar e publicar a evidência até o 60º dia."
    )
    base.add_callout(
        doc, "Prioridade",
        "Primeiro, alcançar 100% dos essenciais; depois, concluir obrigatórios e recomendados. Cada "
        "validação deve testar disponibilidade, atualidade, série histórica, exportação e filtros.",
        fill=base.PALE_GOLD, accent=base.GOLD
    )
    grouped: OrderedDict[str, list[dict[str, str]]] = OrderedDict()
    for row in criteria:
        grouped.setdefault(row["dimensao"], []).append(row)
    for dimension, rows in grouped.items():
        heading = doc.add_heading(dimension, level=2)
        heading.paragraph_format.keep_with_next = True
        if dimension == "Ouvidorias":
            heading.paragraph_format.page_break_before = True
        table = doc.add_table(rows=1, cols=5)
        widths = [620, 1250, 3100, 2868, 1800]
        base.set_table_geometry(table, widths)
        base.repeat_header(table.rows[0])
        base.prevent_row_split(table.rows[0])
        for cell, label in zip(table.rows[0].cells, ["ID", "CLASSE", "CRITÉRIO PNTP", "AÇÃO EM 60 DIAS", "EVIDÊNCIA"]):
            base.set_cell_shading(cell, base.NAVY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(label)
            base.set_run_font(r, 8.0, base.WHITE, bold=True)
        for idx, row in enumerate(rows, 1):
            tr = table.add_row()
            base.prevent_row_split(tr)
            values = [row["id"], row["classe"], row["criterio"], action_for(row), evidence_for(row)]
            for col, (cell, value) in enumerate(zip(tr.cells, values)):
                if idx % 2 == 0:
                    base.set_cell_shading(cell, base.LIGHT_GRAY)
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 0.95
                r = p.add_run(value)
                color = base.BLUE if col == 0 else base.INK
                base.set_run_font(r, 8.0, color, bold=col == 0)
                base.set_cell_border(cell, bottom={"val": "single", "sz": "3", "color": "D8DEE3"})
        note = doc.add_paragraph("Prazo: validar e concluir todos os itens aplicáveis desta dimensão até o 60º dia.")
        note.paragraph_format.space_before = Pt(3)
        note.paragraph_format.space_after = Pt(7)
        for run in note.runs:
            base.set_run_font(run, 8.2, base.MID_GRAY, italic=True)


def add_closing(doc: Document) -> None:
    doc.add_heading("Encaminhamento proposto", level=1)
    base.add_callout(
        doc, "Convite à cooperação",
        "Propõe-se que a Prefeitura acolha este documento como contribuição técnica, designe ponto focal "
        "e publique o cronograma de 60 dias. O objetivo comum é entregar informação pública no prazo, "
        "em formato utilizável e com evidência verificável, fortalecendo a confiança e o controle social.",
        fill=base.PALE_GREEN, accent=base.GREEN
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("FISCALIZA VARGINHA")
    base.set_run_font(r, 11, base.NAVY, bold=True)
    r = p.add_run("\nIniciativa cidadã independente")
    base.set_run_font(r, 9.5, base.MID_GRAY)


def build() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    criteria = load_criteria()
    doc = Document()
    configure_document(doc)
    add_title_block(doc)
    add_main_content(doc, criteria)
    doc.add_page_break()
    add_annex(doc, criteria)
    add_closing(doc)
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build())
